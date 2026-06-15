# -*- coding: utf-8 -*-
"""
FILMN9 밸류에이션 탭 상세 해설서 — Word 파일 생성
삼성전기(009150) 기준 데이터로 각 탭의 모든 값을 해석
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 기본 폰트
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 색상
INDIGO = RGBColor(0x4f, 0x46, 0xe5)
NAVY = RGBColor(0x1b, 0x3a, 0x5b)
GRAY = RGBColor(0x64, 0x74, 0x8b)
ORANGE = RGBColor(0xd9, 0x77, 0x06)
GREEN = RGBColor(0x16, 0xa3, 0x4a)
RED = RGBColor(0xdc, 0x26, 0x26)
PURPLE = RGBColor(0x7c, 0x3a, 0xed)
LIGHT_BG = RGBColor(0xf1, 0xf5, 0xf9)
DARK_TEXT = RGBColor(0x1e, 0x29, 0x3b)

def h(text, level=1, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if level == 0:
        run.font.size = Pt(22)
    elif level == 1:
        run.font.size = Pt(17)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return p

def para(text, bold=False, italic=False, color=None, size=10.5, indent=0):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)
    return p

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"“{text}”")
    run.italic = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(10)
    return p

def bullet_multi(parts):
    """parts: list of (text, bold) tuples"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.size = Pt(10)
        run.bold = bold
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = head
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def callout(title, text, color=INDIGO):
    """강조 박스"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f"💡 {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.3)
    run2 = p2.add_run(text)
    run2.font.size = Pt(10)
    run2.font.name = '맑은 고딕'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p2.paragraph_format.space_after = Pt(8)

def divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(0xc0, 0xc0, 0xc0)
    run.font.size = Pt(9)

# ═════════════════════════════════════════════════════════════════
# 표지
# ═════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FILMN9 밸류에이션 탭')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = INDIGO
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('완전 해설서 — 모든 용어·값·해석법')
run.font.size = Pt(15)
run.font.color.rgb = GRAY
run.italic = True
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('비전공자도 이해할 수 있도록 작성한 정밀 가이드  ·  삼성전기(009150) 데이터 기준')
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()
divider()

# ═════════════════════════════════════════════════════════════════
# 개요
# ═════════════════════════════════════════════════════════════════
h('📋 개요 — 밸류에이션 탭은 무엇을 보여주는가', level=1, color=INDIGO)

para('밸류에이션(Valuation)이란 "이 회사의 1주가 얼마짜리인지" 계산하는 작업입니다. '
     '단 한 가지 방법으로 계산하면 오차가 크기 때문에, 저희 FILMN9 밸류에이션 탭은 '
     '서로 다른 5가지 방법으로 동시에 계산하고 그 근거를 모두 공개합니다.', size=11)
para('')

h('▶ 6개 서브탭 한눈에 보기', level=2, color=INDIGO)
add_table(
    ['탭', '한 줄 설명', '핵심 질문'],
    [
        ['① DCF 분석', '미래에 벌 현금을 모두 합산해 현재가치로 환산', '"이 회사 본질 가치는 얼마?"'],
        ['② WACC', '미래 현금을 깎는 할인율의 산출 근거 공개', '"왜 11.48%로 할인했나?"'],
        ['③ 멀티플', '동종업계가 받는 평균 가격 배수 역산', '"동종업계라면 얼마?"'],
        ['④ 민감도', 'WACC × 성장률 9칸 표로 적정가 변동 확인', '"가정이 틀리면 얼마나 변할까?"'],
        ['⑤ 토네이도', '6개 변수가 적정가에 미치는 영향력 막대그래프', '"무엇을 가장 주시할까?"'],
        ['⑥ 피어 비교', '동종업계 4개사 β·멀티플 비교', '"동종업계와 비교하면?"'],
    ]
)
para('')

callout('이 문서 사용법',
        '각 탭별로 ① 화면에 표시된 모든 값의 정확한 의미, ② 그 값이 적정가 판단에 어떻게 쓰이는지, '
        '③ 초보자가 어떻게 해석해야 하는지 — 3단계로 정리했습니다.',
        color=ORANGE)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# Tab 1: WACC
# ═════════════════════════════════════════════════════════════════
h('💰 [Tab 1] WACC — 미래 돈을 깎는 할인율', level=1, color=NAVY)
para('CAPM 기반 자기자본비용 + 세후 타인자본비용 가중평균', italic=True, color=GRAY, size=10)
para('')

h('▶ 1. WACC가 뭔가요? — 가장 쉬운 비유', level=2, color=INDIGO)
para('"1년 뒤에 받을 100만원과, 지금 받는 100만원은 같은 가치인가?" 답은 NO. '
     '오늘 받으면 은행에 넣어도 이자가 붙기 때문에 미래의 100만원은 오늘 기준으로는 더 적은 가치입니다. '
     '얼마나 적게 봐야 할까? 바로 그 "깎는 비율"이 WACC입니다.', size=10.5)
para('')
quote('미래에 100억을 벌 회사 = WACC 10%로 깎으면 오늘 기준 약 62억(5년 뒤 100억 → 62억). '
      'WACC가 높을수록(=위험할수록) 미래 돈은 더 가혹하게 깎입니다.')
para('')

h('▶ 2. 화면 표시 값 완전 해설 (삼성전기 기준)', level=2, color=INDIGO)

add_table(
    ['항목 (영문)', '한글 풀이', '삼성전기 값', '의미'],
    [
        ['WACC', '가중평균 자본비용', '11.48%', '미래 현금을 11.48% 비율로 깎아서 봄'],
        ['E%', '자기자본 비중', '99.4%', '회사 자금 99.4%가 주주 돈 (부채 거의 없음)'],
        ['D%', '타인자본 비중', '0.6%', '회사 자금 0.6%만 차입금 (안정적)'],
        ['법인세율', '한계 법인세율', '27.5%', '매출 3000억 초과 구간 세율'],
        ['Rf', '무위험수익률 (Risk-free rate)', '4.17%', '국고채 10년물 금리 (안전자산 기준점)'],
        ['β (베타)', '주가 변동성 계수', '0.96', '시장이 1% 움직일 때 0.96% 따라 움직임'],
        ['ERP', '주식위험프리미엄', '8%', '주식이 국고채보다 더 줘야 하는 추가 보상'],
        ['Ke', '자기자본비용 (Cost of Equity)', '11.53%', '주주가 요구하는 수익률'],
        ['Kd', '타인자본비용 (Cost of Debt)', '4.55%', '회사채 발행 시 이자율 (세전)'],
        ['세후 Kd', '세후 타인자본비용', '3.30%', '이자비용이 세금공제 받아 실효 부담은 3.3%'],
    ],
    col_widths=[3.5, 3.5, 2.5, 6]
)
para('')

h('▶ 3. WACC 계산 공식 — 한 줄로 이해하기', level=2, color=INDIGO)
para('WACC = (자기자본비중 × 자기자본비용) + (타인자본비중 × 세후 타인자본비용)', bold=True, size=11)
para('       = (99.4% × 11.53%) + (0.6% × 3.30%)', size=10.5, indent=0.5)
para('       = 11.46% + 0.02% ≈ 11.48%', size=10.5, indent=0.5)
para('')

para('자기자본비용(Ke)은 CAPM(자본자산가격결정모형)으로 계산합니다:', size=10.5)
para('Ke = Rf + β × ERP = 4.17% + 0.96 × 8% = 11.53%', bold=True, size=10.5, indent=0.5)
para('  → 안전자산(국고채) 수익률 위에, 주식 위험만큼 추가 보상을 더한 값', italic=True, color=GRAY, size=9.5)
para('')

h('▶ 4. 각 값이 적정가에 미치는 영향', level=2, color=INDIGO)
add_table(
    ['값이 ↑ 올라가면', 'WACC', '적정주가', '의미'],
    [
        ['Rf (금리 인상)', '↑', '↓', '금리 오르면 주식 매력 떨어져 적정가 하락'],
        ['β (변동성 증가)', '↑', '↓', '주가 변동 클수록 위험 → 더 가혹하게 깎음'],
        ['ERP (시장위험 증가)', '↑', '↓', '시장 전체 위험 증가 시 모든 주식 적정가 하락'],
        ['Kd (회사채 금리)', '↑', '↓', '회사 신용도 악화 → 차입비용 증가 → 적정가 하락'],
        ['세율', '↓', '↑', '세율 낮으면 세후 부담 줄어 적정가 상승'],
    ]
)
para('')

callout('초보자 핵심 메시지',
        'WACC = 8~13%가 일반적입니다. 삼성전기 11.48%는 평균보다 약간 높은 편 (변동성 베타·고금리 환경). '
        'WACC가 1%p 변하면 적정가는 보통 15~25% 변합니다. → 가장 영향 큰 변수입니다!',
        color=RED)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# Tab 2: 멀티플
# ═════════════════════════════════════════════════════════════════
h('📊 [Tab 2] 멀티플 — 동종업계 평균 가격 배수', level=1, color=NAVY)
para('업종 피어그룹 평균 멀티플 적용 역산 적정주가', italic=True, color=GRAY, size=10)
para('')

h('▶ 1. 멀티플이 뭔가요?', level=2, color=INDIGO)
para('"비슷한 회사들이 시장에서 얼마짜리 가격표를 받고 있나?"를 보는 방식입니다. '
     '예를 들어 동종업계 평균 PER이 20배인데, 우리 회사 EPS가 1,000원이면 → 동종업계와 같은 평가를 받으면 주가는 20,000원이어야 한다는 계산입니다.', size=10.5)
para('')

h('▶ 2. 화면 표시 값 완전 해설 (삼성전기 기준)', level=2, color=INDIGO)

add_table(
    ['멀티플 (영문)', '한글 풀이', '피어 평균', '역산 적정가', '상승여력'],
    [
        ['EV/EBITDA', '기업가치 ÷ 영업현금흐름', '44.1배', '₩1,121,196', '-6.9%'],
        ['EV/Sales', '기업가치 ÷ 매출액', '6.8배', '₩1,064,572', '-11.6%'],
        ['PER', '주가 ÷ 주당순이익', '268.4배', '₩2,507,694', '+108.3%'],
        ['PBR', '주가 ÷ 주당순자산', '11.4배', '₩1,498,772', '+24.5%'],
    ]
)
para('* 멀티플 역산 단순평균: ₩1,548,000  (DCF Base ₩84,807 대비 검증용 참고치)', italic=True, color=GRAY, size=9.5)
para('')

h('▶ 3. 각 멀티플의 의미와 사용처', level=2, color=INDIGO)

para('① EV/EBITDA (Enterprise Value to EBITDA)', bold=True, color=PURPLE)
para('  • EV (기업가치) = 시가총액 + 순부채', size=10)
para('  • EBITDA (영업현금흐름) = 영업이익 + 감가상각비', size=10)
para('  • 의미: "회사를 통째로 인수하면 영업현금 몇 년치로 회수되나" — 인수합병·M&A에서 가장 많이 사용', size=10)
para('  • 일반: 8~12배 (낮을수록 저평가, 높을수록 고평가)', size=10)
para('')

para('② EV/Sales', bold=True, color=PURPLE)
para('  • 기업가치 ÷ 매출액', size=10)
para('  • 의미: 매출 1원당 시장이 매기는 가격', size=10)
para('  • 적자 기업 평가나 성장주 평가에 유용 (이익이 음수여도 매출은 있으니까)', size=10)
para('')

para('③ PER (Price to Earnings Ratio)', bold=True, color=PURPLE)
para('  • 주가 ÷ 주당순이익 (EPS)', size=10)
para('  • 의미: "지금 주가로 사면 몇 년 만에 이익으로 회수되나" — 가장 대중적', size=10)
para('  • 일반: 10~20배 (저PER → 저평가 신호 / 고PER → 고성장 기대)', size=10)
para('  • 삼성전기 268.4배는 비정상적으로 높음 → 일시적 이익 급감 영향 (해석 주의)', size=9.5, italic=True, color=RED)
para('')

para('④ PBR (Price to Book Ratio)', bold=True, color=PURPLE)
para('  • 주가 ÷ 주당순자산 (BPS)', size=10)
para('  • 의미: 시장이 회사 청산가치 대비 몇 배 가격을 매기는가', size=10)
para('  • PBR < 1 → 시장이 청산가치보다 싸게 평가 (저평가 신호 또는 망할 위험)', size=10)
para('  • PBR > 2 → 무형자산·브랜드 가치 인정', size=10)
para('')

h('▶ 4. "역산 적정가" 계산 원리', level=2, color=INDIGO)
para('피어 평균 멀티플을 우리 회사 펀더멘털에 곱해서 역산:', size=10.5)
para('예) 피어 평균 PER 22배 × 우리 회사 EPS 1,000원 = 22,000원이 적정가', size=10.5, indent=0.5, italic=True, color=GRAY)
para('')

h('▶ 5. 상승여력 (vs 현재가) 해석', level=2, color=INDIGO)
add_table(
    ['신호', '의미', '해석'],
    [
        ['+ (양수, 초록)', '동종업계 대비 저평가', '동종업계 평균 평가를 받으면 더 오를 여지'],
        ['- (음수, 빨강)', '동종업계 대비 고평가', '동종업계보다 비싸게 거래 중'],
        ['4개 멀티플 일치', '신뢰도 높음', '여러 잣대로 같은 결론 → 확신 있는 적정가'],
        ['4개 멀티플 불일치', '신뢰도 낮음', '사업 특수성 또는 시장 혼란 — 추가 분석 필요'],
    ]
)
para('')

callout('초보자 핵심 메시지',
        'PER이 가장 직관적입니다. 단, 일시적 이익 충격이 있으면(예: 삼성전기 PER 268배) 비정상값이 됩니다. '
        '이럴 때는 EV/EBITDA나 PBR로 교차 검증하세요. 4가지 멀티플의 적정가가 비슷한 범위에 모이면 → 신뢰할 수 있는 적정가입니다.',
        color=GREEN)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# Tab 3: 민감도
# ═════════════════════════════════════════════════════════════════
h('🌡️ [Tab 3] 민감도 (Sensitivity Analysis) — 가정의 견고함 검증', level=1, color=NAVY)
para('WACC (행) × 터미널 성장률 g (열) 조합별 주당 적정가  ·  ◀ = Base case', italic=True, color=GRAY, size=10)
para('')

h('▶ 1. 민감도가 뭔가요?', level=2, color=INDIGO)
para('"내가 입력한 WACC와 성장률이 살짝 틀리면 적정가는 얼마나 변할까?"를 한눈에 보여주는 격자(표)입니다. '
     '밸류에이션은 정답이 아닌 "가정의 묶음"이기 때문에, 가정이 틀렸을 때의 충격을 미리 알아야 합니다.', size=10.5)
para('')

h('▶ 2. 화면 표시 값 완전 해설 (삼성전기 기준)', level=2, color=INDIGO)

add_table(
    ['WACC ↓ / g →', 'g = 1.5%', 'g = 2.0% (Base)', 'g = 2.5%'],
    [
        ['10.5%', '9만', '10만', '10만'],
        ['11.5% ◀ (Base)', '8만', '8만 (Base)', '9만'],
        ['12.5%', '7만', '8만', '8만'],
    ]
)
para('')

h('▶ 3. 용어 풀이', level=2, color=INDIGO)
para('① WACC (가중평균 자본비용)', bold=True)
para('  • 미래 현금을 깎는 비율 — Tab 1에서 11.48%로 산출됨', size=10)
para('  • 민감도에서는 ±1%p 범위에서 변동시켜 봄 (10.5%, 11.5%, 12.5%)', size=10)
para('')
para('② g (영구성장률 / Terminal Growth Rate)', bold=True)
para('  • 5년 예측 이후 회사가 영원히 성장할 비율', size=10)
para('  • 일반적으로 GDP 성장률(약 2~3%) 수준에서 설정', size=10)
para('  • 민감도에서 1.5% / 2.0% / 2.5% 3개 시나리오 비교', size=10)
para('')

h('▶ 4. 표 해석 방법 — 3가지 핵심 패턴', level=2, color=INDIGO)

para('패턴 1: 좌상단(WACC 낮고 g 높음) → 적정가 최대', bold=True, color=GREEN)
para('  → 낙관 시나리오. 금리 인하 + 고성장 지속 시.', size=10, indent=0.3)
para('')

para('패턴 2: 우하단(WACC 높고 g 낮음) → 적정가 최소', bold=True, color=RED)
para('  → 비관 시나리오. 금리 인상 + 저성장 지속 시.', size=10, indent=0.3)
para('')

para('패턴 3: 중앙(◀) Base Case → 가장 가능성 높은 적정가', bold=True, color=INDIGO)
para('  → 현재 가정 기준 적정가 ₩84,807 (8만원대)', size=10, indent=0.3)
para('')

h('▶ 5. 색상 범례 (상승여력 기반)', level=2, color=INDIGO)
add_table(
    ['색상', '의미', '해석'],
    [
        ['진초록', '+50% 초과 상승여력', '극도로 저평가'],
        ['초록', '+10~50% 상승여력', '저평가'],
        ['연초록', '0~+10%', '소폭 저평가'],
        ['노랑', '-10~0%', '거의 적정'],
        ['주황', '-30~-10%', '소폭 고평가'],
        ['빨강', '-30% 미만', '크게 고평가'],
    ]
)
para('')

callout('초보자 핵심 메시지',
        '표 전체가 빨강이면 → "어떤 가정을 써도 현재가가 비싸다" (확실한 고평가 신호). '
        '표 전체가 초록이면 → "어떤 가정을 써도 현재가가 싸다" (확실한 저평가 신호). '
        '색상이 섞이면 → 가정에 매우 민감하다는 의미, 더 신중한 분석 필요.',
        color=ORANGE)

para('')
callout('적정가 판단법',
        '① Base Case의 적정가를 기준선으로 잡는다. '
        '② 표 전체의 최소~최대 범위를 "가능성 있는 적정가 구간"으로 본다. '
        '③ 현재가가 그 구간 안에 있는지, 위인지, 아래인지로 종합 판단한다.',
        color=INDIGO)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# Tab 4: 토네이도
# ═════════════════════════════════════════════════════════════════
h('🌪️ [Tab 4] 토네이도 (Tornado Chart) — 변수별 영향력 순위', level=1, color=NAVY)
para('주요 가정 변동이 적정주가에 미치는 영향 (Base: ₩84,807)', italic=True, color=GRAY, size=10)
para('')

h('▶ 1. 토네이도 차트가 뭔가요?', level=2, color=INDIGO)
para('각 변수를 일정 폭(±2%p 또는 ±1%p)으로 흔들었을 때, 적정주가가 얼마나 출렁이는지 막대그래프로 보여줍니다. '
     '영향력 큰 변수가 위쪽에 위치 (회오리=토네이도 모양에서 이름).', size=10.5)
para('')

h('▶ 2. 화면 표시 값 완전 해설 (삼성전기 기준)', level=2, color=INDIGO)

add_table(
    ['변수 (영문)', '한글 풀이', '하방 (-)', '상방 (+)', '총 변동폭'],
    [
        ['OPM ±2%p', '영업이익률 ±2%p', '-30,685', '+30,685', '61,370 (1위)'],
        ['CapEx율 ±2%p', '자본적지출 비중 ±2%p', '-12,056', '+12,423', '24,479 (2위)'],
        ['WACC ±1%p', '할인율 ±1%p', '-9,410', '+11,728', '21,138 (3위)'],
        ['g (영구성장) ±0.5%p', '터미널 성장률 ±0.5%p', '1,840', '+2,044', '3,884'],
        ['매출 성장률 Y+1 ±2%p', '내년 매출 성장률 ±2%p', '1,707', '+1,770', '3,477'],
        ['NWC율 ±2%p', '순운전자본 비중 ±2%p', '1,566', '+1,566', '3,132'],
    ]
)
para('')

h('▶ 3. 각 변수 의미 상세', level=2, color=INDIGO)

para('① OPM (Operating Profit Margin, 영업이익률)', bold=True, color=PURPLE)
para('  • 매출액 대비 영업이익 비율', size=10)
para('  • ±2%p 변동 시 ±30,685원 (적정가의 36%!) — 가장 큰 영향', size=10)
para('  • 의미: 회사의 본업 수익성이 적정가의 핵심', size=10)
para('')

para('② CapEx (Capital Expenditure, 자본적지출)', bold=True, color=PURPLE)
para('  • 공장·설비·R&D 투자 등 자본 지출 비중', size=10)
para('  • CapEx 증가 → FCFF(잉여현금) 감소 → 적정가 하락', size=10)
para('  • 삼성전기 같은 제조업·반도체 기업에서 특히 중요', size=10)
para('')

para('③ WACC (할인율)', bold=True, color=PURPLE)
para('  • Tab 1에서 다룬 그 WACC. ±1%p 변동 시 ±1만원 수준', size=10)
para('')

para('④ g (Terminal Growth, 영구성장률)', bold=True, color=PURPLE)
para('  • 5년 이후 영원히 성장할 비율 (보통 GDP 수준 2%)', size=10)
para('  • 삼성전기는 영향이 작음 → 5년 이후 성장 가정에 덜 민감한 사업 모델', size=10)
para('')

para('⑤ 매출 성장률 Y+1', bold=True, color=PURPLE)
para('  • 내년(1년 후) 매출 성장률', size=10)
para('  • 단기 매출이 적정가에 미치는 영향 (단기보단 중장기 OPM이 더 중요)', size=10)
para('')

para('⑥ NWC (Net Working Capital, 순운전자본)', bold=True, color=PURPLE)
para('  • 유동자산 - 유동부채 (운영에 묶여있는 돈)', size=10)
para('  • NWC율 증가 → 현금이 재고·매출채권에 묶임 → FCFF 감소', size=10)
para('')

h('▶ 4. 토네이도 차트 해석 방법', level=2, color=INDIGO)

para('① 위에서부터 1~3번 변수를 집중 모니터링', bold=True)
para('  → 삼성전기의 경우: OPM > CapEx > WACC', size=10, indent=0.3)
para('  → 분기 실적에서 영업이익률·설비투자·금리 동향 체크', size=10, indent=0.3)
para('')

para('② 색상으로 변동 방향 확인', bold=True)
para('  • 빨강(하방 변동): 변수가 나쁘게 바뀔 때 적정가 하락', size=10, indent=0.3)
para('  • 초록(상방 변동): 변수가 좋게 바뀔 때 적정가 상승', size=10, indent=0.3)
para('')

para('③ 대칭/비대칭 확인', bold=True)
para('  • OPM 같이 대칭 (-/+ 동일) → 변동이 양방향으로 균등', size=10, indent=0.3)
para('  • WACC처럼 비대칭 (-9,410 vs +11,728) → 한 방향으로 더 민감', size=10, indent=0.3)
para('')

callout('초보자 핵심 메시지',
        '토네이도 차트는 "모니터링 우선순위 가이드"입니다. 모든 뉴스·지표를 다 챙길 시간이 없으니, '
        '상위 3개 변수만 집중하세요. 삼성전기의 경우: 영업이익률(분기실적), 설비투자(자본지출 공시), 금리(WACC). '
        '나머지는 부차적입니다.',
        color=RED)

para('')
callout('적정가 판단법',
        '상위 변수가 "내가 예상한 방향"과 일치하는지 확인하세요. '
        '예: OPM이 개선될 것 같다면 → 적정가 +30,685원 상향 조정. '
        '여러 변수가 같은 방향으로 움직이면 큰 폭의 적정가 변동을 예상해야 합니다.',
        color=INDIGO)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# Tab 5: 피어 비교
# ═════════════════════════════════════════════════════════════════
h('🔍 [Tab 5] 피어 비교 — 동종업계 경쟁사와 나란히 보기', level=1, color=NAVY)
para('업종 피어그룹 4개사 멀티플 비교 · FY2024 기준', italic=True, color=GRAY, size=10)
para('')

h('▶ 1. 피어 비교가 뭔가요?', level=2, color=INDIGO)
para('"우리 회사가 동종업계 경쟁사와 비교해서 비싼가, 싼가, 위험한가, 수익성은 좋은가?"를 한눈에 비교하는 표입니다. '
     '회사 하나만 보면 비싼지 싼지 모르지만, 5개 옆에 놓고 보면 상대 위치가 명확해집니다.', size=10.5)
para('')

h('▶ 2. 화면 표시 값 완전 해설 (삼성전기 기준)', level=2, color=INDIGO)

add_table(
    ['기업', 'β (베타)', 'EV/EBITDA', 'P/E', 'P/B', 'ROE'],
    [
        ['삼성전기 (현재 종목)', '1.2', '—', '—', '—', '—'],
        ['대덕전자', '0.96', '—', '—', '—', '—'],
        ['두산', '1.25', '—', '—', '—', '—'],
        ['리노공업', '1.43', '—', '—', '—', '—'],
        ['피어 평균', '1.21', '44.1x', '6.8x', '268.4x', '—'],
    ]
)
para('')

h('▶ 3. 각 지표 의미', level=2, color=INDIGO)

para('① β (Beta, 베타) — 변동성 위험도', bold=True, color=PURPLE)
para('  • 시장(코스피)이 1% 움직일 때 이 주식이 몇 % 움직이는지', size=10)
para('  • β > 1: 시장보다 출렁임 큼 (변동성 높음, 공격적)', size=10)
para('  • β < 1: 시장보다 안정적', size=10)
para('  • 삼성전기 1.2 → 시장보다 20% 더 변동적 (반도체 사이클 영향)', size=10)
para('')

para('② EV/EBITDA — 영업현금흐름 배수', bold=True, color=PURPLE)
para('  • 기업가치 ÷ 영업현금흐름', size=10)
para('  • 낮을수록 저평가, 높을수록 고평가', size=10)
para('  • 피어 평균 44.1배는 비정상적으로 높음 → 일시적 EBITDA 하락 영향', size=10)
para('')

para('③ P/E (Price to Earnings) — 주당이익 배수', bold=True, color=PURPLE)
para('  • 주가 ÷ 주당순이익', size=10)
para('  • 가장 대중적인 지표. 저PER = 저평가 신호 (단, 일시 충격 주의)', size=10)
para('')

para('④ P/B (Price to Book) — 주당순자산 배수', bold=True, color=PURPLE)
para('  • 주가 ÷ 주당순자산 (BPS)', size=10)
para('  • PBR < 1 → 시장이 청산가치보다 싸게 평가', size=10)
para('  • PBR > 2 → 무형자산·브랜드·기술력 인정', size=10)
para('')

para('⑤ ROE (Return on Equity) — 자기자본수익률', bold=True, color=PURPLE)
para('  • 순이익 ÷ 자기자본 × 100%', size=10)
para('  • 회사가 주주 돈으로 얼마나 효율적으로 이익을 내는지', size=10)
para('  • ROE > 10%: 우량 / ROE > 15%: 매우 우수', size=10)
para('')

h('▶ 4. 피어 비교 해석 매트릭스', level=2, color=INDIGO)

add_table(
    ['상황', '판단', '액션'],
    [
        ['우리 PER < 평균 + ROE > 평균', '저평가 + 우량', '🟢 매수 우선 검토'],
        ['우리 PER < 평균 + ROE < 평균', '저평가 + 실적 부진', '🟡 가치주 함정 주의'],
        ['우리 PER > 평균 + ROE > 평균', '고평가 + 우량', '🟡 프리미엄 종목, 추격매수 신중'],
        ['우리 PER > 평균 + ROE < 평균', '고평가 + 부진', '🔴 매도 우선 검토'],
        ['우리 β > 평균', '동종업계보다 변동성 큼', '리스크 관리 필수'],
    ]
)
para('')

h('▶ 5. 피어 그룹 선정의 중요성', level=2, color=INDIGO)
para('삼성전기 피어: 대덕전자, 두산, 리노공업 (전자부품·인쇄회로기판 업종)', size=10)
para('• 같은 사업영역의 기업을 선정해야 의미 있음', size=10)
para('• 시가총액 규모가 비슷한 기업 우선 (대형주는 대형주끼리)', size=10)
para('• 해외 글로벌 피어 포함 시 시장 평가 추가 검증 가능', size=10)
para('')

callout('초보자 핵심 메시지',
        '단일 지표가 아니라 "조합"으로 봐야 합니다. PER이 낮다고 무조건 저평가가 아니라, ROE가 좋고 β가 낮으면 진짜 저평가입니다. '
        'PER 낮은데 ROE도 낮으면 "가치주 함정"일 수 있습니다 (실적이 나빠서 주가가 싼 것).',
        color=GREEN)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# 종합 적정가 판단법
# ═════════════════════════════════════════════════════════════════
h('🎯 종합 — 5개 탭을 묶어 적정가 결정하는 4단계', level=1, color=INDIGO)
para('')

h('▶ Step 1. 적정가 범위 확정', level=2, color=INDIGO)
para('각 방법론에서 나온 적정가를 모은다:', size=10.5)
add_table(
    ['방법론', '적정가', '신뢰도'],
    [
        ['DCF Base', '₩84,807', '⭐⭐⭐ 중심값'],
        ['민감도 표 최소', '₩70,000대', '시나리오 하한'],
        ['민감도 표 최대', '₩90,000대', '시나리오 상한'],
        ['멀티플 평균', '₩1,548,000', '⚠️ 비정상 (피어 멀티플 왜곡)'],
    ]
)
para('→ 신뢰 가능한 적정가 범위: 약 ₩70,000 ~ ₩90,000 (DCF 기반)', bold=True, color=INDIGO)
para('')

h('▶ Step 2. 핵심 변수 점검 (토네이도 활용)', level=2, color=INDIGO)
para('상위 3개 변수가 현재 어떤 방향으로 움직이고 있는지 확인:', size=10.5)
bullet('OPM (영업이익률): 분기 실적·가이던스 체크')
bullet('CapEx (설비투자): 공시 자본지출 계획 체크')
bullet('WACC (금리): 한국은행 기준금리·국고채 금리 추이')
para('')

h('▶ Step 3. 동종업계 위치 확인 (피어 비교)', level=2, color=INDIGO)
para('우리 종목이 동종업계 대비 어디 위치하는지:', size=10.5)
bullet('β 비교: 변동성이 동종업계 평균보다 큰가, 작은가')
bullet('수익성(ROE): 동종업계보다 자본효율이 좋은가')
bullet('밸류에이션(PER/PBR): 동종업계 평균보다 저평가/고평가')
para('')

h('▶ Step 4. 종합 판단 — 신호등 매트릭스', level=2, color=INDIGO)
add_table(
    ['상황', '신호', '액션'],
    [
        ['현재가 < 민감도 표 최소', '🟢 강력 저평가', '비중 확대 검토'],
        ['민감도 표 최소 ≤ 현재가 ≤ Base', '🟢 저평가', '분할 매수'],
        ['Base < 현재가 ≤ 민감도 표 최대', '🟡 약간 고평가', '관망 또는 분할 매수'],
        ['현재가 > 민감도 표 최대', '🔴 고평가', '비중 축소 또는 매도'],
    ]
)
para('')

h('▶ 삼성전기 종합 판단 (현재가 ₩1,340,000 기준)', level=2, color=INDIGO)
para('• DCF 적정가 ₩84,807 ↔ 현재가 ₩1,340,000', size=10.5)
para('• 현재가가 DCF 적정가의 약 15배 → 🔴 극도의 고평가 신호', size=10.5, color=RED)
para('• 단, PER/PBR/EV/EBITDA 멀티플로는 -7% ~ -41% 수준', size=10.5)
para('• 해석: 시장은 미래 성장(AI 서버용 MLCC, 전장 사업)을 매우 낙관적으로 반영 중', size=10.5)
para('  → DCF는 현재 펀더멘털만 보고 있어 보수적, 멀티플은 시장 기대를 반영', size=10, italic=True, color=GRAY)
para('')

callout('최종 메시지',
        '하나의 숫자에 매달리지 마세요. DCF는 보수적 본질가치, 멀티플은 시장 평가, 민감도는 가정의 견고함, '
        '토네이도는 리스크 우선순위, 피어 비교는 상대 위치 — 5가지를 종합해서 "감"을 잡는 것이 가장 합리적입니다.',
        color=INDIGO)

para('')
divider()

# 부록
h('📌 부록: 자주 묻는 질문', level=1, color=INDIGO)

para('Q1. 왜 DCF 적정가와 멀티플 적정가가 이렇게 차이 나나요?', bold=True)
para('A. DCF는 "현재 펀더멘털 기반 본질가치", 멀티플은 "시장이 현재 매기는 가격". '
     '시장이 미래 성장을 강하게 기대할 때(AI/전장)는 멀티플이 DCF를 크게 상회합니다.', size=10)
para('')

para('Q2. WACC 11.48%가 적당한 수치인가요?', bold=True)
para('A. 한국 코스피 대형주 평균 WACC가 8~12% 수준이므로 약간 높은 편입니다. '
     '베타 0.96과 무위험금리 4.17%(고금리 환경) 영향이 큽니다.', size=10)
para('')

para('Q3. 멀티플 PER 268배는 어떻게 봐야 하나요?', bold=True)
para('A. 일시적 이익 급감(특별손실, 일회성 비용 등)으로 분모가 작아진 영향입니다. '
     '이럴 때는 PER보다 EV/EBITDA·EV/Sales를 우선 참고하세요.', size=10)
para('')

para('Q4. 민감도 표가 모두 빨강인데 절대 사면 안 되는 건가요?', bold=True)
para('A. "DCF 모델 기준" 고평가일 뿐, 시장은 미래 성장을 반영합니다. '
     '추가로 멀티플·피어 비교·시장 모멘텀까지 종합 판단해야 합니다.', size=10)
para('')

para('Q5. 적정가가 시장가와 크게 다르면 모델이 틀린 건가요?', bold=True)
para('A. 둘 다 정답이 아닙니다. 시장가는 "현재 시장 참여자들의 평균 의견", '
     '적정가는 "재무 데이터 기반 모델 추정값". 둘의 괴리가 클수록 → 시장이 강한 기대(또는 우려)를 반영 중. '
     '괴리의 방향과 크기를 해석하는 것이 투자자의 일입니다.', size=10)
para('')

divider()

para('© FILMN9 · KPMG AI Lab · 2026-05-26', italic=True, color=GRAY, size=9)
para('본 문서는 학습·참고용이며, 특정 종목의 매수·매도를 권유하지 않습니다.', italic=True, color=GRAY, size=9)

# 저장
output_path = r"C:\Users\Admin\FILMN9\FILMN9_밸류에이션_완전해설서.docx"
doc.save(output_path)
print(f"저장 완료: {output_path}")
