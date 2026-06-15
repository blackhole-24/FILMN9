# -*- coding: utf-8 -*-
"""
FILMN9 데모 발표 스크립트 — 3개 탭 워드 파일 생성
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 기본 폰트 설정 (맑은 고딕)
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_heading(text, level=1, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if level == 0:
        run.font.size = Pt(20)
        if color: run.font.color.rgb = color
    elif level == 1:
        run.font.size = Pt(16)
        if color: run.font.color.rgb = color
    elif level == 2:
        run.font.size = Pt(13)
        if color: run.font.color.rgb = color
    else:
        run.font.size = Pt(11)
        if color: run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p

def add_para(text, bold=False, italic=False, color=None, indent=False, size=10.5):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    return p

def add_quote(text):
    """발표 멘트 (이탤릭, 들여쓰기)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"\"{text}\"")
    run.italic = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1f, 0x3a, 0x5f)
    return p

def add_bullet(text, level=0, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    if bold_part:
        # bold_part는 (앞, 굵게, 뒤) 튜플
        for i, segment in enumerate(bold_part):
            run = p.add_run(segment)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.size = Pt(10)
            run.bold = (i == 1)
    else:
        run = p.add_run(text)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.size = Pt(10)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    # 데이터
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return table

def add_divider():
    p = doc.add_paragraph()
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(0xc0, 0xc0, 0xc0)
    run.font.size = Pt(8)

# ─────────────────────────────────────────────────────────────────
# 표지
# ─────────────────────────────────────────────────────────────────
INDIGO = RGBColor(0x4f, 0x46, 0xe5)
NAVY = RGBColor(0x1b, 0x3a, 0x5b)
GRAY = RGBColor(0x64, 0x74, 0x8b)
ORANGE = RGBColor(0xd9, 0x77, 0x06)
GREEN = RGBColor(0x16, 0xa3, 0x4a)
RED = RGBColor(0xdc, 0x26, 0x26)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FILMN9 데모 발표 스크립트')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = INDIGO
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI 기반 기업 분석 자동화 플랫폼  ·  3개 탭 시연')
run.font.size = Pt(13)
run.font.color.rgb = GRAY
run.italic = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('KPMG AI Lab  ·  2026-05-26  ·  소요시간 약 6분')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_paragraph()
add_divider()

# ─────────────────────────────────────────────────────────────────
# 오프닝
# ─────────────────────────────────────────────────────────────────
add_heading('🎯 오프닝 (0:00–0:30)', level=1, color=INDIGO)
add_quote("주식 살 때 가장 어려운 질문, '이 회사 지금 어떤 상황이고, 가격은 적정한가?' 입니다.")
add_quote("애널리스트 보고서는 30페이지, 초보자에겐 너무 어렵죠. 토스 증권에 들어가도 숫자만 잔뜩 나옵니다.")
add_quote("저희 FILMN9은 — 한 종목에 대해 기업분석·밸류에이션·AI상담을 5분 만에 끝낼 수 있도록 3개 탭으로 정리했습니다.")
add_para("")
add_para("→ 3종목 데모: 아모레퍼시픽(090430) · 삼성전기(009150) · NAVER(035420)", bold=True, color=INDIGO)
add_para("")

# ─────────────────────────────────────────────────────────────────
# 탭 1: 기업개요
# ─────────────────────────────────────────────────────────────────
add_heading('📊 [Tab 1] 기업개요 — 5분 안에 회사 파악 (0:30–2:30)', level=1, color=NAVY)

add_heading('▶ 핵심 메시지', level=2, color=INDIGO)
add_quote("기업개요 탭은 '이 회사 처음 본 사람이 5분 안에 회사 전체를 파악할 수 있게' 만든 종합 대시보드입니다.")
add_para("")

add_heading('▶ 화면 구성 & 발표 포인트', level=2, color=INDIGO)

add_para("① Hero 배너 + 실시간 헤더", bold=True)
add_bullet("종목코드 · 회사명 · KOSPI 배지 · 업종(WICS) · 시총 · 대표이사 한눈에")
add_bullet("실시간 주가 (Yahoo Finance · 60초 폴링 · ~15분 지연 · 장중 표시)")
add_quote("KRX 실시간 데이터피드는 연간 수천만원 유료이지만, 저희는 yfinance API로 15분 지연 데이터를 60초마다 자동 갱신합니다.")

add_para("② 히스토리 브리핑 (LLM 자동 요약)", bold=True)
add_bullet("DART 사업보고서를 LLM(gpt-5-mini)이 자동 요약 · 최신 뉴스(Google News RSS) 6건")
add_bullet("주가 영향 요인 + 사업 모델 + 리스크 요약")
add_quote("DART 사업보고서 200페이지를 LLM이 핵심만 추려서 보여줍니다. 신뢰도 높음(high) 등급이면 그대로 활용 가능합니다.")

add_para("③ 주가차트 (lightweight-charts)", bold=True)
add_bullet("기본 1Y · 기간 전환(1M/3M/6M/1Y/3Y/전체) · MA5/20/60 + 볼린저밴드 토글")
add_bullet("차트 내 휠 = 줌인/줌아웃 · Y축 위 휠 = 페이지 스크롤 (UX 개선)")
add_bullet("[i] 아이콘 → 실시간 차트 한계 사유 안내 (KRX WebSocket 유료 등)")

add_para("④ 재무 하이라이트 + 신용등급 추이", bold=True)
add_bullet("최근 매출액 · 영업이익 · 당기순이익 + YoY 색상(빨강↑/파랑↓)")
add_bullet("신용등급 5년 추이 라인차트 (AAA → BBB 매핑)")

add_para("⑤ 기업개요 + 경영인 (2컬럼)", bold=True)
add_bullet("DART 사업보고서: 시총·상장일·임직원·연혁·주요제품 + 경영진 리스트")

add_para("⑥ 주주구성 + 기업건전성 (2컬럼)", bold=True)
add_bullet(None, bold_part=("주주구성 도넛차트 — ", "최대주주 5인 + 기타", " 비중"))
add_bullet(None, bold_part=("기업건전성 3종 — ", "부채비율 · 유동비율 · 영업이익률", " (B/S 자동 계산)"))
add_bullet("등급 신호등: 🟢 우량 / 🟡 주의 / 🔴 위험 + 종합 등급")
add_quote("부채비율·유동비율은 DART 사업보고서의 B/S 데이터를 자체 계산식으로 산출합니다. 시중 서비스에서 안 보여주는 지표까지 자동 계산하는 게 차별점입니다.")

add_para("⑦ 재무제표 탭 (BS / IS / Sankey)", bold=True)
add_bullet("Balance Sheet · Income Statement 5년치 + 출처 라벨")
add_bullet("Sankey 흐름도 — 매출 → 영업비용 → 순이익까지 시각화")

add_para("⑧ 최근 공시 10건", bold=True)
add_bullet("DART OpenAPI 실시간 연동 · 클릭 시 원문 PDF 이동")
add_para("")

add_heading('▶ 핵심 차별점 한 줄', level=2, color=INDIGO)
add_quote("네이버 증권은 숫자만 나열, 저희는 'AI가 요약 + 자동 계산 + 시각화'까지 5분 패키지로 제공합니다.")
add_para("")

# ─────────────────────────────────────────────────────────────────
# 탭 2: 밸류에이션
# ─────────────────────────────────────────────────────────────────
add_heading('💎 [Tab 2] 밸류에이션 — 적정주가 다각도 검증 (2:30–4:30)', level=1, color=NAVY)

add_heading('▶ 핵심 메시지', level=2, color=INDIGO)
add_quote("'이 주식 지금 사도 되는 가격인가?' — 5가지 방법으로 동시에 계산하고, 산출 근거까지 한 화면에 보여줍니다.")
add_para("")

add_heading('▶ 상단 3카드 — 결론 (Bear/Base/Bull)', level=2, color=INDIGO)
add_table(
    ['카드', '의미', '초보자 활용법'],
    [
        ['🐻 BEAR', '최악 시나리오 적정가', '바닥 가격대 (안전마진)'],
        ['📊 BASE', '가장 가능성 높은 적정가', '저평가 / 고평가 1차 판단'],
        ['🐂 BULL', '최상 시나리오 도달가', '목표가 / 차익실현 기준'],
    ]
)
add_quote("단일 가격이 아닌 '가격의 범위'를 제시 — 현재가가 어디 위치하는지 확률적으로 판단합니다.")
add_para("")

add_heading('▶ 하위 6탭 — 적정가 산출 근거', level=2, color=INDIGO)
add_table(
    ['탭', '핵심', '한 줄 해석'],
    [
        ['📈 DCF', '미래 현금흐름 → 현재가치', '"부동산처럼 5년 월세+매각가 합산한 본질가치"'],
        ['💰 WACC', '할인율 산출 근거', '"왜 8.2%로 할인했는지 투명 공개"'],
        ['📊 멀티플', 'PER·PBR·EV/EBITDA 시장 비교', '"4개 지표 적정가가 비슷하면 신뢰도 ↑"'],
        ['🌡️ 민감도', 'WACC × 성장률 표', '"내 판단의 신뢰 구간"'],
        ['🌪️ 토네이도', '변수별 영향력 순위', '"무엇을 최우선 모니터링할지 가이드"'],
        ['🔍 피어 비교', '국내+글로벌 경쟁사 비교', '"동종업계 대비 비싼가 싼가, 수익성은?"'],
    ]
)
add_para("")

add_heading('▶ 페르소나 1 (초보자) 활용 시나리오 — 3분 안에 의사결정', level=2, color=INDIGO)
add_bullet(None, bold_part=("① 카드 3장 → ", "\"현재가가 BASE보다 높네, 고평가\"", ""))
add_bullet(None, bold_part=("② DCF·멀티플 → ", "\"다각도로 봐도 결과 일치, 신뢰도 ↑\"", ""))
add_bullet(None, bold_part=("③ 토네이도 → ", "\"WACC가 핵심 변수, 금리 뉴스 챙기자\"", ""))
add_bullet(None, bold_part=("④ 피어 비교 → ", "\"동종업계 대비론 저평가\"", ""))
add_quote("→ 분할 매수 전략 결정. 30페이지 애널리스트 보고서를 3분 만에 다각도 검증.")
add_para("")

add_heading('▶ 핵심 차별점 한 줄', level=2, color=INDIGO)
add_quote("네이버 증권은 PER 1개, 저희는 '15개 적정가 + 산출 근거 + 리스크 가이드'를 한 화면에.")
add_para("")
add_para("※ 참고: 향후 가치진단 5종 (EPV·성장 프리미엄·시장함의 성장률·기대 격차·안전마진) 추가 예정 — 밸류에이션 파트 산출물 통합 시 업그레이드", italic=True, color=GRAY, size=9)
add_para("")

# ─────────────────────────────────────────────────────────────────
# 탭 3: AI분석
# ─────────────────────────────────────────────────────────────────
add_heading('🤖 [Tab 3] AI분석 — 사업보고서 대화형 검색 (4:30–5:30)', level=1, color=NAVY)

add_heading('▶ 핵심 메시지', level=2, color=INDIGO)
add_quote("사업보고서 200페이지를 다 읽을 시간이 없으신가요? AI에게 직접 물어보세요. 출처까지 알려드립니다.")
add_para("")

add_heading('▶ 화면 구성', level=2, color=INDIGO)
add_para("① AI 종합분석 (상단)", bold=True)
add_bullet("강점 / 약점 / 기회 / 위협 4가지 관점 자동 요약")
add_bullet("DART 사업보고서 + 재무 데이터를 LLM이 종합 분석")

add_para("② RAG 챗봇 (하단)", bold=True)
add_bullet("종목별 추천 질문 칩 3개 제공 (클릭 시 자동 답변)")
add_bullet("자유 입력 시 키워드 매칭 → 사업보고서 기반 답변")
add_bullet("답변 하단에 출처 표기 → 신뢰도 확보")
add_para("")

add_heading('▶ 종목별 추천 질문 (시연 시 자동 표시)', level=2, color=INDIGO)

add_para("🌸 아모레퍼시픽 (090430)", bold=True, color=RGBColor(0xd9, 0x77, 0x06))
add_bullet("주요 브랜드 포트폴리오는?")
add_bullet("중국 의존도와 미국·일본 시장 성과는?")
add_bullet("R&D 투자와 기술 경쟁력은?")

add_para("🟦 삼성전기 (009150)", bold=True, color=RGBColor(0x25, 0x63, 0xeb))
add_bullet("MLCC 시장 지위와 경쟁력은?")
add_bullet("전장(자동차) 사업 성장 전망은?")
add_bullet("삼성전자 매출 의존도는?")

add_para("🟪 NAVER (035420)", bold=True, color=RGBColor(0x7c, 0x3a, 0xed))
add_bullet("주요 사업 부문과 매출 비중은?")
add_bullet("AI(하이퍼클로바X)·클라우드 전략은?")
add_bullet("주요 규제·경쟁 리스크는?")
add_para("")

add_heading('▶ 실제 답변 예시 (삼성전기 MLCC 질문)', level=2, color=INDIGO)
add_quote("삼성전기 주력 사업 MLCC는 글로벌 시장 점유율 약 23%로 일본 무라타(35%)에 이은 세계 2위입니다. AI 서버·전기차 수요 확대로 24년 매출 +18%. 자동차용 고신뢰성 MLCC가 ASP 3~5배 높은 고부가 영역입니다. (출처: DART 2024 사업보고서 II.사업의 내용)")
add_para("")

add_heading('▶ 핵심 차별점 한 줄', level=2, color=INDIGO)
add_quote("ChatGPT는 일반 지식, 저희는 'DART 공시 + 출처 표기' RAG로 신뢰도 확보. 투자 의사결정에 직접 활용 가능합니다.")
add_para("")

# ─────────────────────────────────────────────────────────────────
# 클로징
# ─────────────────────────────────────────────────────────────────
add_heading('🏆 클로징 (5:30–6:00)', level=1, color=INDIGO)
add_para("")

add_heading('▶ 3개 탭 = 5분 패키지', level=2, color=INDIGO)
add_table(
    ['탭', '핵심 가치', '소요시간'],
    [
        ['Tab 1 기업개요', '5분 만에 회사 전체 파악', '~2분'],
        ['Tab 2 밸류에이션', '15개 적정가 + 다각도 검증', '~2분'],
        ['Tab 3 AI분석', '사업보고서 대화형 검색', '~1분'],
    ]
)
add_para("")

add_heading('▶ 서비스 타당성 — 시장 비교', level=2, color=INDIGO)
add_table(
    ['항목', '기존 서비스', 'FILMN9'],
    [
        ['기업 정보', '숫자만 나열', 'AI 자동 요약 + 5분 패키지'],
        ['적정주가', '단일 컨센서스 1개', '3시나리오 × 5방법론 = 15개'],
        ['리스크 분석', '없음', '민감도 + 토네이도 + 가치진단'],
        ['사업보고서', '직접 PDF 읽기', 'RAG 챗봇으로 5초 답변'],
        ['초보자 대응', '용어만 어려움', '직관 시각화 + 우선순위 가이드'],
    ]
)
add_para("")

add_heading('▶ 마무리 한 문장', level=2, color=INDIGO)
add_quote("기관 애널리스트가 30페이지 보고서로 하는 일을, 비전공자도 5분 만에. 이게 FILMN9이 존재하는 이유입니다.")
add_quote("숫자를 보여주는 서비스는 많지만, 숫자의 이유까지 보여주는 서비스는 우리뿐입니다.")
add_para("")

add_divider()

# ─────────────────────────────────────────────────────────────────
# 발표 체크리스트 (부록)
# ─────────────────────────────────────────────────────────────────
add_heading('📌 발표 체크리스트', level=1, color=INDIGO)
add_bullet("Tab 1: 'AI 자동 요약 + 자체 계산 지표(부채비율·유동비율)' 강조")
add_bullet("Tab 2: '가격이 아닌 가격의 범위' + '다각도 교차 검증' 강조")
add_bullet("Tab 3: '출처까지 표기되는 RAG' — ChatGPT와의 차별점 강조")
add_bullet("종목 전환 시 추천 질문 칩이 자동으로 바뀌는 것 보여주기")
add_bullet("Yahoo Finance 15분 지연 시각이 헤더에 실시간 표시되는 것 자랑")
add_bullet("마무리는 '숫자의 이유까지 보여주는 유일한 서비스' 멘트로 임팩트")
add_para("")

# 저장
output_path = r"C:\Users\Admin\FILMN9\FILMN9_데모_발표스크립트_3탭.docx"
doc.save(output_path)
print(f"저장 완료: {output_path}")
