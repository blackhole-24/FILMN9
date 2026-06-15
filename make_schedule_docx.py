# -*- coding: utf-8 -*-
"""FILMN9 정밀 스케줄 워드 파일 생성"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 기본 폰트
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.font.size = Pt(10)

# 여백
for section in doc.sections:
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(text, size=18, color=(56, 189, 248), bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    return p

def add_para(text, size=10, color=(60, 60, 60), bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    return p

def set_cell_text(cell, text, size=9, bold=False, color=(40, 40, 40), align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# ===== 표지 =====
add_heading('FILMN9 정밀 스케줄', size=22, color=(15, 76, 117))
add_heading('2026.05.29(금) → 2026.06.05(금)', size=13, color=(100, 100, 100), bold=False)
add_para('스프린트 4 마무리 · 6/4(목) 오전 9시 컷오프 → AWS 통합 → 6/5(금) 시연 발표',
         size=10, color=(80, 80, 80))
add_para('작성일: 2026-05-29(금) · 작성자: 기업개요 파트', size=9, color=(120, 120, 120))
doc.add_paragraph()

# ===== 범례 =====
add_heading('범례', size=12, color=(15, 76, 117))
add_para('A 컬럼 = 오늘 새로 지시한 업무  /  B 컬럼 = 이전 스케줄에서 아직 못 끝낸 업무',
         size=10)
add_para('⏱ = 예상 소요 시간', size=10)
doc.add_paragraph()

# ===== 일자별 스케줄 =====
DAYS = [
    {
        "title": "🔴 5/29 (금) — 오늘 · D-7",
        "header_color": "EF4444",
        "subtotal": "6.5h",
        "sections": [
            {
                "header": None,
                "rows": [
                    ("1", "A", "Supabase 통합 타당성 사전 조사", "SQLite 200.7MB + MongoDB 2,526docs 용량 측정", "1.0h"),
                    ("2", "A", "재무정보 검증 프로그램 v1 설계", "Task 1~14 자동 일일 리포트 골조", "2.0h"),
                    ("3", "A", "주가차트 자동 동기화 요구사항 정리", "pykrx/yfinance · 일배치 시간대 결정", "0.5h"),
                    ("4", "B", "Task 15 스키마 수정", "current_price → ohlcv 종가 매핑", "0.5h"),
                    ("5", "B", "연결<별도 55개 종목 파싱 재검증", "CFS/OFS 라벨 뒤바뀐 원인 추적", "2.0h"),
                    ("6", "B", "백엔드 상시 가동 + Task 10·13 재실행", "어제 결측 API 성능·데모 검증 재측정", "0.5h"),
                ],
            },
        ],
    },
    {
        "title": "🟠 5/30 (토) · D-6",
        "header_color": "F59E0B",
        "subtotal": "9.0h",
        "sections": [
            {
                "header": None,
                "rows": [
                    ("1", "A", "재무정보 검증 프로그램 v1 코딩", "auto_full_analysis.py → daily_finance_validator.py", "3.0h"),
                    ("2", "A", "주가차트 OHLCV 자동 동기화 cron 설계", "매일 16:30 종가 적재 스케줄러", "1.5h"),
                    ("3", "A", "Sankey 디자인 이슈 정리", "3종목 캡처 분석 · 색상·라벨·노드 정렬", "1.0h"),
                    ("4", "B", "외감법인 422개 JSONL 청크 파싱 스크립트 골조", "parse_jsonl_chunks_to_db.py 작성", "3.0h"),
                    ("5", "B", "ui_payload 통합 사전 협의 (밸류에이션 파트 DM)", "스키마·필드 정합 확인", "0.5h"),
                ],
            },
        ],
    },
    {
        "title": "🟡 5/31 (일) · D-5",
        "header_color": "FBBF24",
        "subtotal": "11.0h",
        "sections": [
            {
                "header": None,
                "rows": [
                    ("1", "A", "재무정보 3000개 기업 1차 전수 검증 실행", "현재 1,861개 → 외감 422 합산 후 검증", "2.0h"),
                    ("2", "A", "주가차트 자동 동기화 구현", "pykrx → ohlcv 일배치 + 누락일 백필", "3.0h"),
                    ("3", "A", "Sankey 차트 디자인 수정 1차", "Plotly 컬러 팔레트 통일 · 라벨 한글화", "2.0h"),
                    ("4", "B", "외감법인 422개 DB 적재 실행 + 무결성 검증", "financial_detail INSERT 후 PK 중복 확인", "3.0h"),
                    ("5", "B", "삼성전기 PER 268배 비정상 수치 재검증", "분모 EPS 산출 로직 추적", "1.0h"),
                ],
            },
        ],
    },
    {
        "title": "🟢 6/1 (월) · D-4",
        "header_color": "10B981",
        "subtotal": "10.5h",
        "sections": [
            {
                "header": None,
                "rows": [
                    ("1", "A", "재무정보 '의미있는 값' 도출 로직", "성장률·영업이익률·ROE·부채비율 자동 계산 컬럼", "4.0h"),
                    ("2", "A", "히스토리브리핑 검증 프로그램 v1", "기업개요 파트 MongoDB 2,526docs 무결성·중복·NULL 검증", "3.0h"),
                    ("3", "A", "Supabase 타당성 보고서 작성", "SQLite+Mongo 통합 vs 유지 / 마이그 비용 / RLS 장점", "1.5h"),
                    ("4", "B", "WICS 업종 분류 데이터 수집", "KIND/KRX 매핑 → company_info.sector", "2.0h"),
                ],
            },
        ],
    },
    {
        "title": "🔵 6/2 (화) · D-3",
        "header_color": "3B82F6",
        "subtotal": "10.0h",
        "sections": [
            {
                "header": None,
                "rows": [
                    ("1", "A", "메인 홈페이지 검색란 + 모닝루틴 위젯", "전일 시황·등락 TOP5·테마 이슈 설계 + API", "3.0h"),
                    ("2", "A", "산업별 검색 → 업종별 대장주 리스트 API", "/api/sector/{wics}/leaders 작성", "2.0h"),
                    ("3", "A", "Tab3 AI 종합분석 프롬프트 v1", "SWOT·투자포인트·리스크 요약 템플릿", "2.0h"),
                    ("4", "B", "ui_payload JSON → valuations 테이블 적재", "밸류에이션 파트 산출물 통합", "2.0h"),
                    ("5", "B", "DCF 누락 860개 제외 필터 검증", "매출액·자산총계 보유 종목만 DCF 통과", "1.0h"),
                ],
            },
        ],
    },
    {
        "title": "🟣 6/3 (수) · D-2",
        "header_color": "8B5CF6",
        "subtotal": "11.0h",
        "sections": [
            {
                "header": None,
                "rows": [
                    ("1", "A", "AI 사업보고서 RAG 챗봇 정확도 개선", "3종목 → 30종목 확장 · 인용·페이지 표기", "4.0h"),
                    ("2", "A", "Tab3 AI 종합분석 → 프론트 연결", "스트리밍 응답·로딩 UI", "2.0h"),
                    ("3", "A", "메인 홈페이지 산업별 대장주 화면 적용", "v0.dev 컴포넌트 → 실데이터 연결", "2.0h"),
                    ("4", "B", "차별성 4종 자동 추출 1차", "감사의견·리스크 (단일판매·주가영향은 스프린트 5)", "2.0h"),
                    ("5", "B", "재무·히스토리 검증 프로그램 통합 리포트", "두 결과 합산 → HTML 대시보드", "1.0h"),
                ],
            },
        ],
    },
    {
        "title": "⚫ 6/4 (목) — 오전 9시 컷오프 · D-1",
        "header_color": "1F2937",
        "subtotal": "12.0h",
        "sections": [
            {
                "header": "☀️ 오전 (00:00 ~ 09:00) — 내 파트 최종 완성",
                "rows": [
                    ("1", "A", "검증 3종(재무·히스토리·밸류) 전수 실행 → PASS 확인", "", "2.0h"),
                    ("2", "A", "코드 동결 + git tag v0.4-sprint4", "", "0.5h"),
                    ("3", "B", "Sankey 최종 수정·차트 동결", "", "1.0h"),
                ],
            },
            {
                "header": "🌙 오전 9시 이후 — AWS 통합 (형주·밸류에이션 파트·기업개요 파트·소희 합류)",
                "rows": [
                    ("4", "C", "AWS 통합 킥오프", "", "0.5h"),
                    ("5", "C", "IAM · VPC · 보안그룹 세팅", "", "1.5h"),
                    ("6", "C", "RDS(PostgreSQL or Supabase) 스키마 적재", "", "2.0h"),
                    ("7", "C", "S3 임베딩 20GB 업로드", "", "1.5h"),
                    ("8", "C", "EC2 백엔드 배포 + Vercel 프론트 연결", "", "2.0h"),
                    ("9", "C", "통합 리허설 1회", "", "1.0h"),
                ],
            },
        ],
    },
    {
        "title": "🎯 6/5 (금) — D-day · 2차 배포 시연",
        "header_color": "059669",
        "subtotal": "발표",
        "sections": [
            {
                "header": None,
                "rows": [
                    ("1", "T", "09:00 — AWS 환경 헬스체크 + 데모 3종목 사전 검증", "", "1.0h"),
                    ("2", "T", "10:00 — 백업 시연 영상 재생 테스트", "", "1.0h"),
                    ("3", "T", "14:00 — 🎤 2차 배포 시연 발표", "", "2.0h"),
                    ("4", "T", "16:00 — 회고 + 스프린트 5 킥오프", "", "1.0h"),
                ],
            },
        ],
    },
]

for day in DAYS:
    # 일자 제목
    p = doc.add_paragraph()
    run = p.add_run(f"{day['title']}     [소계 {day['subtotal']}]")
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(day['header_color'])

    for section in day['sections']:
        if section['header']:
            sp = doc.add_paragraph()
            srun = sp.add_run(section['header'])
            srun.font.name = '맑은 고딕'
            srun._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            srun.font.size = Pt(10)
            srun.font.bold = True
            srun.font.color.rgb = RGBColor(100, 100, 100)

        # 표
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        widths = [Cm(1.0), Cm(4.5), Cm(6.5), Cm(5.5), Cm(1.5)]
        hdr = table.rows[0].cells
        headers = ["#", "A. 새 지시 업무", "B. 이전 미완료", "상세", "⏱"]
        for i, h in enumerate(headers):
            set_cell_text(hdr[i], h, size=9, bold=True, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_bg(hdr[i], "1F4E79")
            hdr[i].width = widths[i]

        for row in section['rows']:
            num, kind, title, desc, time = row
            tr = table.add_row().cells
            set_cell_text(tr[0], num, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100,100,100))
            if kind == "A":
                set_cell_text(tr[1], title, size=9, bold=True, color=(8, 112, 184))
                set_cell_text(tr[2], "—", size=9, color=(180,180,180), align=WD_ALIGN_PARAGRAPH.CENTER)
            elif kind == "B":
                set_cell_text(tr[1], "—", size=9, color=(180,180,180), align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(tr[2], title, size=9, bold=True, color=(180, 90, 0))
            elif kind == "C":
                set_cell_text(tr[1], title, size=9, bold=True, color=(60, 60, 60))
                set_cell_text(tr[2], "(AWS 통합)", size=9, color=(120,120,120), align=WD_ALIGN_PARAGRAPH.CENTER)
            elif kind == "T":
                set_cell_text(tr[1], title, size=9, bold=True, color=(5, 150, 105))
                set_cell_text(tr[2], "(발표일)", size=9, color=(120,120,120), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(tr[3], desc, size=9, color=(80, 80, 80))
            set_cell_text(tr[4], time, size=9, bold=True, color=(40, 40, 40), align=WD_ALIGN_PARAGRAPH.CENTER)
            for i, w in enumerate(widths):
                tr[i].width = w

        doc.add_paragraph()  # 간격

# ===== 누적 작업량 =====
doc.add_page_break()
add_heading('📊 누적 작업량 (혼자 처리분)', size=14, color=(15, 76, 117))

wl_table = doc.add_table(rows=1, cols=3)
wl_table.style = 'Light Grid Accent 1'
wl_hdr = wl_table.rows[0].cells
for i, h in enumerate(["날짜", "막대 그래프", "시간"]):
    set_cell_text(wl_hdr[i], h, size=10, bold=True, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(wl_hdr[i], "1F4E79")

workload = [
    ("5/29 (금)", 6.5),
    ("5/30 (토)", 9.0),
    ("5/31 (일)", 11.0),
    ("6/1 (월)", 10.5),
    ("6/2 (화)", 10.0),
    ("6/3 (수)", 11.0),
    ("6/4 (목)", 12.0),
]
for date, hrs in workload:
    row = wl_table.add_row().cells
    set_cell_text(row[0], date, size=10, bold=True)
    bar = "█" * int(hrs * 2) + "░" * (24 - int(hrs * 2))
    set_cell_text(row[1], bar, size=9, color=(56, 189, 248))
    set_cell_text(row[2], f"{hrs}h", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# 합계 행
row = wl_table.add_row().cells
set_cell_text(row[0], "합계", size=11, bold=True, color=(15, 76, 117))
set_cell_text(row[1], "7일간 (일 평균 10h)", size=10, color=(15, 76, 117))
set_cell_text(row[2], "70.0h", size=11, bold=True, color=(15, 76, 117), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# ===== 결정 필요 항목 =====
add_heading('⚠️ 결정 필요 항목', size=14, color=(180, 90, 0))

dec_table = doc.add_table(rows=1, cols=3)
dec_table.style = 'Light Grid Accent 2'
dec_hdr = dec_table.rows[0].cells
for i, h in enumerate(["항목", "영향", "결정 시점"]):
    set_cell_text(dec_hdr[i], h, size=10, bold=True, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(dec_hdr[i], "78350F")

decisions = [
    ("Supabase 통합 여부", "6/2(화) RDS 작업 직결 · SQLite+Mongo → Supabase 단일화 검토", "6/1(월)"),
    ("OpenAI gpt-5-mini 사용", "RAG 30종목 확장 비용 ($50 잔액 공유)", "6/2(화)"),
    ("외감법인 422개 보충 우선순위", "데이터 커버리지 1,861 → 2,283 종목 확대", "5/30(토) 시작 전"),
]
for item, impact, when in decisions:
    row = dec_table.add_row().cells
    set_cell_text(row[0], item, size=10, bold=True, color=(100, 50, 0))
    set_cell_text(row[1], impact, size=9, color=(80, 80, 80))
    set_cell_text(row[2], when, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(180, 90, 0))

doc.add_paragraph()

# ===== 핵심 마일스톤 =====
add_heading('🎯 핵심 마일스톤 (4단계)', size=14, color=(15, 76, 117))

ms_table = doc.add_table(rows=1, cols=3)
ms_table.style = 'Light Grid Accent 1'
ms_hdr = ms_table.rows[0].cells
for i, h in enumerate(["단계", "기간", "핵심 산출물"]):
    set_cell_text(ms_hdr[i], h, size=10, bold=True, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(ms_hdr[i], "1F4E79")

milestones = [
    ("① 데이터 정확도", "5/29(금) ~ 5/31(일)", "외감 422개 보충 / 연결<별도 55개 수정 / 재무 검증 v1 / OHLCV 동기화"),
    ("② 의미값·메인", "6/1(월) ~ 6/2(화)", "재무 지표 자동 계산 / 모닝루틴 / 산업별 대장주 / 히스토리 검증"),
    ("③ AI 기능", "6/3(수)", "RAG 30종목 확장 / Tab3 종합분석 / 차별성 추출 1차"),
    ("④ AWS 통합·시연", "6/4(목) ~ 6/5(금)", "코드 동결 / IAM·VPC / RDS·S3·EC2 배포 / 시연 발표"),
]
for stage, period, output in milestones:
    row = ms_table.add_row().cells
    set_cell_text(row[0], stage, size=10, bold=True, color=(15, 76, 117))
    set_cell_text(row[1], period, size=9, color=(80, 80, 80), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row[2], output, size=9, color=(80, 80, 80))

doc.add_paragraph()
add_para('© FILMN9 · 작성일 2026-05-29(금) · 스프린트 4 최종 스케줄',
         size=9, color=(150, 150, 150), align=WD_ALIGN_PARAGRAPH.CENTER)

output = r"C:\Users\Admin\FILMN9\FILMN9_스케줄_5월29일_6월5일.docx"
doc.save(output)
print(f"[OK] {output}")
