# -*- coding: utf-8 -*-
"""
마크다운 → Word(.docx) 변환기 v2
한글 폰트(동아시아) 속성 완전 지원
"""

import re, os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_KO   = '맑은 고딕'
FONT_CODE = 'Consolas'

# ──────────────────────────────────────────────────────────────
# 유틸: 런에 한글 폰트 완전 설정 (eastAsia 포함)
# ──────────────────────────────────────────────────────────────
def set_run_font(run, name=FONT_KO, size=10.5,
                 bold=False, italic=False, color=None):
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color

    rPr    = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), name)


def set_doc_default_font(doc, name=FONT_KO):
    """문서 기본 스타일에 동아시아 폰트 설정"""
    for style_name in ('Normal', 'Default Paragraph Font'):
        try:
            st = doc.styles[style_name]
            st.font.name = name
            rPr    = st.element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                rFonts.set(qn(attr), name)
        except Exception:
            pass


# ──────────────────────────────────────────────────────────────
# 셀 배경색
# ──────────────────────────────────────────────────────────────
def hex6(color: RGBColor) -> str:
    return str(color)   # RGBColor.__str__ → '1E40AF'

def set_cell_bg(cell, color: RGBColor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6(color))
    tcPr.append(shd)


# ──────────────────────────────────────────────────────────────
# 단락 테두리 (하단선)
# ──────────────────────────────────────────────────────────────
def add_bottom_border(para, color='1E40AF'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ──────────────────────────────────────────────────────────────
# 단락 왼쪽 세로선 (인용구)
# ──────────────────────────────────────────────────────────────
def add_left_border(para, color='1E40AF'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    lft  = OxmlElement('w:left')
    lft.set(qn('w:val'),   'single')
    lft.set(qn('w:sz'),    '12')
    lft.set(qn('w:space'), '4')
    lft.set(qn('w:color'), color)
    pBdr.append(lft)
    pPr.append(pBdr)


# ──────────────────────────────────────────────────────────────
# 인라인 마크다운 파싱 (볼드 / 이탤릭 / 인라인코드)
# ──────────────────────────────────────────────────────────────
INLINE_RE = re.compile(
    r'(\*\*\*(.+?)\*\*\*'   # ***bold+italic***
    r'|\*\*(.+?)\*\*'       # **bold**
    r'|\*(.+?)\*'           # *italic*
    r'|`([^`]+?)`)'         # `code`
)

def parse_inline(para, text: str, size=10.5, color=None):
    last = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > last:
            run = para.add_run(text[last:m.start()])
            set_run_font(run, size=size, color=color)
        g = m.groups()
        if g[1]:   # bold+italic
            run = para.add_run(g[1])
            set_run_font(run, size=size, bold=True, italic=True, color=color)
        elif g[2]: # bold
            run = para.add_run(g[2])
            set_run_font(run, size=size, bold=True, color=color)
        elif g[3]: # italic
            run = para.add_run(g[3])
            set_run_font(run, size=size, italic=True, color=color)
        elif g[4]: # inline code
            run = para.add_run(g[4])
            set_run_font(run, name=FONT_CODE, size=9.5)
        last = m.end()
    if last < len(text):
        run = para.add_run(text[last:])
        set_run_font(run, size=size, color=color)


# ──────────────────────────────────────────────────────────────
# 헤더 추가
# ──────────────────────────────────────────────────────────────
C_H1  = RGBColor(0x1E, 0x40, 0xAF)
C_H2  = RGBColor(0x1D, 0x4E, 0x89)
C_H3  = RGBColor(0x1E, 0x40, 0xAF)
C_H4  = RGBColor(0x37, 0x51, 0x8C)
C_WHT = RGBColor(0xFF, 0xFF, 0xFF)
C_TBL = RGBColor(0x1E, 0x40, 0xAF)
C_ALT = RGBColor(0xEF, 0xF6, 0xFF)
C_QUO = RGBColor(0x37, 0x51, 0x8C)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=20, bold=True, color=C_H1)
    add_bottom_border(p, '1E40AF')
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=15, bold=True, color=C_H2)
    add_bottom_border(p, '1D4E89')
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=12.5, bold=True, color=C_H3)
    return p

def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run('▶ ' + text)
    set_run_font(run, size=11, bold=True, color=C_H4)
    return p


# ──────────────────────────────────────────────────────────────
# 일반 단락
# ──────────────────────────────────────────────────────────────
def add_normal(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.6)
    parse_inline(p, text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.4 + level * 0.5)
    # 수동 불릿 기호
    run0 = p.add_run('• ')
    set_run_font(run0, size=10.5, bold=True)
    parse_inline(p, text)
    return p


def add_numbered(doc, text, num):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5)
    parse_inline(p, f'{num}. ' + text)
    return p


# ──────────────────────────────────────────────────────────────
# 코드 블록
# ──────────────────────────────────────────────────────────────
def add_code_block(doc, lines):
    for raw in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        # 회색 배경
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F3F4F6')
        pPr.append(shd)
        run = p.add_run(raw if raw else ' ')
        set_run_font(run, name=FONT_CODE, size=9)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


# ──────────────────────────────────────────────────────────────
# 인용구 (blockquote)
# ──────────────────────────────────────────────────────────────
def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    add_left_border(p, '1E40AF')
    parse_inline(p, text, size=10.5, color=C_QUO)
    for run in p.runs:
        run.italic = True
    return p


# ──────────────────────────────────────────────────────────────
# 수평선
# ──────────────────────────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    add_bottom_border(p, 'D1D5DB')


# ──────────────────────────────────────────────────────────────
# 마크다운 테이블 → Word 테이블
# ──────────────────────────────────────────────────────────────
def parse_table(lines, i):
    header_cols = [c.strip() for c in lines[i].strip().split('|') if c.strip()]
    j = i + 1
    # 구분선 스킵
    while j < len(lines) and re.match(r'^\s*\|[\s\-|:]+\|\s*$', lines[j]):
        j += 1
    data_rows = []
    while j < len(lines) and '|' in lines[j]:
        row = [c.strip() for c in lines[j].strip().split('|') if c.strip() != '']
        if row:
            data_rows.append(row)
        j += 1
    return header_cols, data_rows, j


def add_md_table(doc, header, rows):
    ncols = len(header)
    if ncols == 0:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 헤더 행
    hcells = tbl.rows[0].cells
    for ci, h in enumerate(header):
        hcells[ci].text = ''
        p = hcells[ci].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(h)
        set_run_font(run, size=9.5, bold=True, color=C_WHT)
        set_cell_bg(hcells[ci], C_TBL)

    # 데이터 행
    for ri, row in enumerate(rows):
        bg = C_ALT if ri % 2 == 1 else None
        cells = tbl.rows[ri + 1].cells
        for ci in range(ncols):
            val = row[ci] if ci < len(row) else ''
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            parse_inline(p, val, size=9.5)
            if bg:
                set_cell_bg(cells[ci], bg)

    doc.add_paragraph()   # 테이블 뒤 공백


# ──────────────────────────────────────────────────────────────
# 핵심 변환 함수
# ──────────────────────────────────────────────────────────────
def convert(md_path: str, docx_path: str):
    doc = Document()
    set_doc_default_font(doc, FONT_KO)

    # 페이지 여백
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    i          = 0
    in_code    = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # ── 코드 블록 토글 ─────────────────────────────────────
        if line.strip().startswith('```'):
            if not in_code:
                in_code    = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code    = False
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # ── 빈 줄 ──────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── 수평선 ─────────────────────────────────────────────
        if re.match(r'^[-=]{3,}\s*$', stripped):
            add_hr(doc)
            i += 1
            continue

        # ── 헤더 ───────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            lvl  = len(m.group(1))
            text = m.group(2).strip()
            [add_h1, add_h2, add_h3, add_h4][lvl - 1](doc, text)
            i += 1
            continue

        # ── 마크다운 테이블 ────────────────────────────────────
        if stripped.startswith('|') and i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if re.match(r'^\|[\s\-|:]+\|$', nxt):
                hdr, rows, ni = parse_table(lines, i)
                add_md_table(doc, hdr, rows)
                i = ni
                continue

        # ── 인용구 ─────────────────────────────────────────────
        if stripped.startswith('>'):
            add_blockquote(doc, stripped.lstrip('> '))
            i += 1
            continue

        # ── 체크박스 ───────────────────────────────────────────
        m = re.match(r'^(\s*)- \[([ xX])\]\s+(.*)', line)
        if m:
            chk    = m.group(2).lower() == 'x'
            lvl    = len(m.group(1)) // 2
            prefix = '☑ ' if chk else '☐ '
            add_bullet(doc, prefix + m.group(3), level=lvl)
            i += 1
            continue

        # ── 순서 없는 목록 ─────────────────────────────────────
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            lvl = len(m.group(1)) // 2
            add_bullet(doc, m.group(2), level=lvl)
            i += 1
            continue

        # ── 순서 있는 목록 ─────────────────────────────────────
        m = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if m:
            add_numbered(doc, m.group(3), m.group(2))
            i += 1
            continue

        # ── 일반 텍스트 ────────────────────────────────────────
        add_normal(doc, stripped)
        i += 1

    doc.save(docx_path)
    fname = os.path.basename(docx_path)
    sz    = os.path.getsize(docx_path)
    print(f'  [OK] {fname}  ({sz:,} bytes)')


# ──────────────────────────────────────────────────────────────
# 실행
# ──────────────────────────────────────────────────────────────
DOCS_DIR = r'C:\Users\Admin\FILMN9\docs'
OUT_DIR  = r'C:\Users\Admin\FILMN9\docs\word'

FILES = [
    ('01_FILMN9_프로젝트기획서_v1.md',  '01_FILMN9_프로젝트기획서_v1.docx'),
    ('02_FILMN9_페르소나_v1.md',        '02_FILMN9_페르소나_v1.docx'),
    ('03_FILMN9_프로젝트헌장_v1.md',    '03_FILMN9_프로젝트헌장_v1.docx'),
    ('04_FILMN9_설계도_워터폴_v1.md',   '04_FILMN9_설계도_워터폴_v1.docx'),
    ('05_FILMN9_설계도_에자일_v1.md',   '05_FILMN9_설계도_에자일_v1.docx'),
]

if __name__ == '__main__':
    sys.stdout = __import__('io').TextIOWrapper(
        sys.stdout.buffer, encoding='utf-8', errors='replace'
    )
    os.makedirs(OUT_DIR, exist_ok=True)
    print(f'[출력 폴더] {OUT_DIR}\n')
    for md_name, docx_name in FILES:
        print(f'  변환 중: {md_name}')
        convert(
            os.path.join(DOCS_DIR, md_name),
            os.path.join(OUT_DIR,  docx_name)
        )
    print('\n[완료] 전체 변환 완료!')
