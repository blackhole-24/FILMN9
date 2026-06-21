# -*- coding: utf-8 -*-
"""
FINSIGHT 용어사전(백과사전) 생성기 — 2026-06-17
금융·회계·투자 + 개발·기술·인프라 용어를 '전문 정의 + 쉬운 비유 + 우리 서비스 활용' 3단으로.
대상: 심사위원·비전공자·초등학생도 한눈에 이해 + 전문가도 인정할 정확성.
출력: 통합산출물/FINSIGHT_용어사전_백과사전.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLUE = RGBColor(0x25, 0x63, 0xEB)
GRAY = RGBColor(0x55, 0x5F, 0x6B)
GREEN = RGBColor(0x15, 0x7A, 0x3C)

def set_kfont(run, name="맑은 고딕"):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name); rFonts.set(qn('w:hAnsi'), name); rFonts.set(qn('w:eastAsia'), name)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor)
    tcPr.append(sh)

doc = Document()
# 기본 폰트
st = doc.styles['Normal']; st.font.name='맑은 고딕'; st.font.size=Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.2); s.right_margin=Cm(2.2)

def H(text, size=20, color=NAVY, before=10, after=6, align=None):
    p = doc.add_paragraph(); p.space_before=Pt(before); p.space_after=Pt(after)
    if align is not None: p.alignment = align
    r = p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color; set_kfont(r)
    return p

def body(text, size=10, color=None, italic=False, after=4):
    p = doc.add_paragraph(); p.space_after=Pt(after)
    r = p.add_run(text); r.font.size=Pt(size); r.italic=italic
    if color: r.font.color.rgb=color
    set_kfont(r); return p

# ── 표지 ──
H("FINSIGHT", 34, NAVY, 60, 2, WD_ALIGN_PARAGRAPH.CENTER)
H("금융·개발 용어 백과사전", 22, BLUE, 2, 4, WD_ALIGN_PARAGRAPH.CENTER)
body("전문가의 정확성과 초보자의 눈높이를 동시에 — 심사위원·비전공자·입문자를 위한 통합 용어집",
     11, GRAY, True, 4).alignment = WD_ALIGN_PARAGRAPH.CENTER
body("운영사 FILMN9 Inc.  ·  2026-06-17  ·  KPMG AI Lab", 10, GRAY, False, 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
body("📖 읽는 법 — 각 용어는 세 줄로 구성됩니다.", 11, NAVY)
body("   · 정의 : 정확한 전문적 의미 (회계사·개발자 기준)", 10)
body("   · 쉽게 : 비전공자·초등학생도 이해하는 일상 비유", 10, GREEN)
body("   · 활용 : FINSIGHT 서비스에서 실제로 쓰이는 위치", 10, BLUE)
doc.add_page_break()

def term(ko, en, definition, easy, usage):
    p = doc.add_paragraph(); p.space_before=Pt(7); p.space_after=Pt(1)
    r = p.add_run(f"▸ {ko}"); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=NAVY; set_kfont(r)
    if en:
        r2 = p.add_run(f"  ({en})"); r2.font.size=Pt(9); r2.font.color.rgb=GRAY; r2.italic=True; set_kfont(r2)
    def line(label, txt, color):
        pp = doc.add_paragraph(); pp.space_after=Pt(1); pp.paragraph_format.left_indent=Cm(0.5)
        rl = pp.add_run(f"{label} "); rl.bold=True; rl.font.size=Pt(9.5); rl.font.color.rgb=color; set_kfont(rl)
        rt = pp.add_run(txt); rt.font.size=Pt(9.5); set_kfont(rt)
    line("정의", definition, GRAY)
    line("쉽게", easy, GREEN)
    line("활용", usage, BLUE)

def section(title):
    doc.add_paragraph()
    p = doc.add_paragraph(); p.space_before=Pt(8); p.space_after=Pt(6)
    r = p.add_run(title); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); set_kfont(r)
    # 배경 음영
    pPr = p._p.get_or_add_pPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'1F3A5F'); pPr.append(sh)

def formula_box(name, lines, variables, note=None):
    """크게 강조한 계산식 박스 + 변수 1:1 풀이."""
    p = doc.add_paragraph(); p.space_before=Pt(10); p.space_after=Pt(2)
    r = p.add_run(f"◆ {name}"); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=NAVY; set_kfont(r)
    # 공식 박스 (단일 셀 표, 음영+테두리)
    tbl = doc.add_table(rows=1, cols=1); tbl.style='Table Grid'
    cell = tbl.rows[0].cells[0]; shade(cell, 'EAF1FB')
    cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before=Pt(8); cp.paragraph_format.space_after=Pt(8)
    for i, ln in enumerate(lines):
        if i>0: cp.add_run().add_break()
        fr = cp.add_run(ln); fr.bold=True; fr.font.size=Pt(15 if len(lines)==1 else 13)
        fr.font.color.rgb=NAVY; set_kfont(fr)
    # 변수 풀이 헤더
    ph = doc.add_paragraph(); ph.space_before=Pt(3); ph.space_after=Pt(1); ph.paragraph_format.left_indent=Cm(0.3)
    rh = ph.add_run("▷ 각 값의 뜻"); rh.bold=True; rh.font.size=Pt(9.5); rh.font.color.rgb=GREEN; set_kfont(rh)
    for sym, mean in variables:
        pp = doc.add_paragraph(); pp.space_after=Pt(1); pp.paragraph_format.left_indent=Cm(0.8)
        rs = pp.add_run(f"{sym}"); rs.bold=True; rs.font.size=Pt(10); rs.font.color.rgb=BLUE; set_kfont(rs)
        rm = pp.add_run(f"  :  {mean}"); rm.font.size=Pt(10); set_kfont(rm)
    if note:
        pn = doc.add_paragraph(); pn.space_after=Pt(4); pn.paragraph_format.left_indent=Cm(0.8)
        rn = pn.add_run(f"※ {note}"); rn.italic=True; rn.font.size=Pt(9.5); rn.font.color.rgb=GRAY; set_kfont(rn)

# ════════════════════════════════════════════════════════
# PART 1 — 금융·회계·투자
# ════════════════════════════════════════════════════════
H("PART 1.  금융 · 회계 · 투자 용어", 18, NAVY, 4, 8)

section("★ 밸류에이션 핵심 계산식 총정리 (반드시 이해할 6대 공식)")
body("FINSIGHT 밸류에이션의 모든 적정주가는 아래 공식들로 계산됩니다. 각 공식과 그 안에 들어가는 값의 뜻을 하나씩 풀이합니다. "
     "(임의 가정·근거 없는 숫자는 쓰지 않는 NO-MOCK 원칙을 따릅니다.)", 10, GRAY, True, 6)

formula_box("① DCF — 기업가치 (절대가치평가의 정석)",
    ["기업가치(EV) = Σ  FCFF(t) ÷ (1 + WACC)^t   +   터미널밸류 ÷ (1 + WACC)^n",
     "적정주가 = ( 기업가치 − 순부채 ) ÷ 발행주식수"],
    [("FCFF(t)", "t년차 잉여현금흐름 — 회사가 자유롭게 쓸 수 있는 현금"),
     ("WACC", "가중평균자본비용 = 할인율 — 미래의 돈을 현재가치로 깎는 비율"),
     ("t", "연차 (1년, 2년 … n년)"),
     ("n", "명시적 예측기간 — 보통 5~10년"),
     ("터미널밸류", "예측기간 이후 회사가 영구히 창출할 가치 (③ 참조)"),
     ("순부채", "총차입금 − 보유현금"),
     ("발행주식수", "시장에 발행된 총 주식 수")],
    "미래 현금흐름을 모두 현재가치로 환산해 더한 것이 기업가치. 채권자 몫(순부채)을 빼면 주주가치, 주식수로 나누면 한 주의 적정주가.")

formula_box("② FCFF — 잉여현금흐름 (DCF의 입력값)",
    ["FCFF = EBIT × (1 − Tc) + D&A − CapEx − ΔNWC"],
    [("EBIT", "이자·세금 차감 전 영업이익 — 본업으로 번 돈"),
     ("Tc", "법인세율 — 이익에 매기는 세금 비율"),
     ("D&A", "감가상각비 — 현금이 실제로 나가지 않는 비용이라 다시 더함(+)"),
     ("CapEx", "자본적지출 — 설비·자산에 쓴 실제 현금이라 뺌(−)"),
     ("ΔNWC", "순운전자본 증감 — 재고·외상에 묶이는 돈의 증가분이라 뺌(−)")],
    "'장부상 이익'이 아니라 '진짜 남는 현금'을 구하는 식. 애널리스트 리포트로 역산·보정해 정밀도를 높인다.")

formula_box("③ 터미널밸류 — 영구가치 (고든 성장모형)",
    ["터미널밸류 = FCFF(n+1) ÷ (WACC − g) = FCFF(n) × (1 + g) ÷ (WACC − g)"],
    [("FCFF(n+1)", "예측 마지막 해 다음 연도의 잉여현금흐름"),
     ("g", "영구성장률 — 장기적으로 유지할 성장률 (보통 1~3%, 보수적으로 설정)"),
     ("WACC", "할인율 (②④에서 산출)")],
    "조건: 반드시 WACC > g (아니면 값이 무한대로 발산). DCF 기업가치의 큰 비중을 차지하므로 g를 과대 설정하지 않는다.")

formula_box("④ WACC — 가중평균자본비용 (DCF의 할인율)",
    ["WACC = (E ÷ V) × Re + (D ÷ V) × Rd × (1 − Tc)"],
    [("E", "자기자본 — 주주 몫(시가총액 기준)"),
     ("D", "타인자본 — 부채(차입금)"),
     ("V", "총자본 = E + D"),
     ("Re", "자기자본비용 — 주주가 요구하는 수익률(⑤ CAPM으로 계산)"),
     ("Rd", "타인자본비용 — 빌린 돈의 이자율"),
     ("Tc", "법인세율 — 이자비용의 절세효과를 반영(1−Tc)")],
    "주주 몫과 빚 몫의 비용을 자본 비중대로 섞은 '종합 자본 단가'. 회사가 돈을 끌어오는 평균 이자율.")

formula_box("⑤ CAPM — 자기자본비용 (Re 산출)",
    ["Re = Rf + β × (Rm − Rf)"],
    [("Re", "자기자본비용 — 구하려는 값(주주의 요구수익률)"),
     ("Rf", "무위험수익률 — 국고채 금리(떼일 걱정 없는 기본 이자)"),
     ("β", "베타 — 이 주식이 시장보다 얼마나 출렁이는지의 민감도"),
     ("Rm", "시장기대수익률 — 주식시장 전체의 평균 기대수익"),
     ("(Rm − Rf)", "시장위험프리미엄(ERP) — 위험을 감수한 대가로 더 받는 수익률")],
    "'안전한 이자(Rf) + 위험에 대한 보상(β×ERP)'으로 주주의 요구수익률을 계산. 베타는 피어 기반 추정 후 Blume·Hamada로 보정.")

formula_box("⑥ 상대가치 역산 — 4-Way 교차검증(멀티플)",
    ["PER 역산 적정주가 = 적정 PER × EPS",
     "PBR 역산 적정주가 = 적정 PBR × BPS",
     "EV/EBITDA :  EV = 적정배수 × EBITDA  →  적정주가 = (EV − 순부채) ÷ 발행주식수"],
    [("적정 PER/PBR/배수", "피어(유사기업)의 중앙값 등으로 정한 기준 배수"),
     ("EPS", "주당순이익 = 당기순이익 ÷ 발행주식수"),
     ("BPS", "주당순자산 = 순자산 ÷ 발행주식수"),
     ("EBITDA", "영업이익 + 감가상각비 — 현금창출력에 근접한 지표")],
    "DCF(절대가치)와 함께 PER·PBR·EV/EBITDA 3가지를 더해 총 4가지 방법으로 교차검증 → 한쪽 편향을 막는다.")

section("1-1. 기업가치 평가의 기본")
term("밸류에이션", "Valuation",
     "기업의 경제적 가치를 정량적으로 추정하는 행위. 미래 현금흐름·자산·이익을 근거로 '이 회사는 얼마짜리인가'를 계산한다.",
     "집을 사기 전 '이 집 적정 가격이 얼마지?'를 따져보는 것과 같다. 회사판 감정평가.",
     "FINSIGHT 밸류에이션 파트의 핵심. 약 2,223개 종목에 자동으로 적정주가를 산출한다.")
term("내재가치", "Intrinsic Value",
     "시장 가격(주가)과 무관하게, 기업이 본질적으로 창출하는 가치. 미래 현금흐름의 현재가치 합으로 추정한다.",
     "물건의 '진짜 값어치'. 세일 가격표(시장가)와 별개로 원래 가치가 얼마인지를 보는 것.",
     "DCF로 계산한 내재가치를 현재 주가와 비교해 고평가/저평가를 진단한다.")
term("적정주가", "Fair Value / Target Price",
     "여러 가치평가 모델로 산출한, 한 주당 '합당한' 가격. 현재 주가와의 차이가 투자 판단의 근거가 된다.",
     "시험 점수의 '정답'에 해당. 지금 시장 가격이 정답보다 싸면 매력적, 비싸면 부담.",
     "종목 상세 밸류에이션 탭에서 적정주가 범위(비관·기본·낙관)로 제시한다.")
term("상승여력", "Upside Potential",
     "현재 주가가 적정주가까지 오를 수 있는 여지. (적정주가 ÷ 현재주가 − 1) × 100(%).",
     "현재 1만원인데 적정가가 1만 2천원이면 '20% 더 오를 여지'가 있다는 뜻.",
     "밸류 스코어카드와 요약 랭킹에서 상승여력으로 종목을 정렬·비교한다.")
term("안전마진", "Margin of Safety",
     "적정가치 대비 시장가격이 얼마나 할인되어 있는지의 폭. 추정 오차를 흡수하는 완충 장치(벤저민 그레이엄 개념).",
     "비올 것 같을 때 우산을 미리 챙기는 여유. 계산이 조금 틀려도 손해 보지 않을 안전 쿠션.",
     "가치진단 지표 중 하나로 표시해 보수적 투자 관점을 제공한다.")
term("절대가치평가 vs 상대가치평가", "Absolute vs Relative Valuation",
     "절대가치평가는 기업 자체의 현금흐름으로 가치를 구하고(DCF), 상대가치평가는 유사기업과 비교해 구한다(멀티플).",
     "절대평가는 '이 음식의 원가로 값 매기기', 상대평가는 '옆 가게보다 비싼지 싼지로 값 매기기'.",
     "FINSIGHT는 둘을 함께 써서 한쪽에 치우치지 않게 교차검증한다.")

section("1-2. 현금흐름 할인법 (DCF)")
term("DCF (현금흐름 할인법)", "Discounted Cash Flow",
     "기업이 미래에 벌어들일 현금흐름을 적정 할인율로 현재가치화하여 합산, 기업가치를 구하는 절대가치평가의 정석.",
     "앞으로 매년 받을 용돈을 '지금 받는다면 얼마일까'로 환산해 모두 더하는 것.",
     "밸류에이션 엔진의 중심 모델. 3개 시나리오로 적정주가 범위를 만든다.")
term("FCF / FCFF (잉여현금흐름)", "Free Cash Flow to Firm",
     "FCFF = EBIT×(1−법인세율) + 감가상각비(D&A) − 자본적지출(CapEx) − 순운전자본증감(ΔNWC). 채권자·주주 모두에게 귀속되는 현금.",
     "회사가 운영비·투자비를 다 쓰고 '자유롭게 남긴 현금'. 진짜 곳간에 쌓이는 돈.",
     "DCF의 입력값. 애널리스트 리포트로 FCFF를 역산·보정해 정밀도를 높인다.")
term("현재가치 / 할인율", "Present Value / Discount Rate",
     "미래의 돈을 현재 시점 가치로 환산한 값이 현재가치, 그 환산에 쓰는 비율이 할인율(보통 WACC).",
     "1년 뒤 100만원은 지금 100만원보다 가치가 낮다(이자·위험 때문). 그 차이를 반영하는 비율이 할인율.",
     "FCFF를 WACC로 할인해 기업가치를 구한다.")
term("영구가치 (터미널 밸류)", "Terminal Value",
     "명시적 예측기간(보통 5~10년) 이후 기업이 영구히 창출할 가치. TV = FCFF₍ₙ₊₁₎ ÷ (WACC − g).",
     "10년치만 계산하고 '그 뒤로도 회사는 계속 돈 번다'는 부분을 한 덩어리로 묶은 값.",
     "DCF 기업가치의 큰 비중을 차지하므로 영구성장률(g)을 보수적으로 둔다.")
term("영구성장률", "Perpetual Growth Rate (g)",
     "영구가치 계산 시 가정하는, 기업이 장기적으로 유지할 성장률. 보통 물가상승률·GDP 성장률 수준(1~3%).",
     "회사가 늙어도 '매년 조금씩은 더 큰다'고 가정하는 아주 완만한 성장 속도.",
     "과대 가정 시 가치가 폭증하므로 NO-MOCK 원칙에 따라 보수적으로 설정한다.")
term("시나리오 분석 (Bear/Base/Bull)", "Scenario Analysis",
     "비관(Bear)·기본(Base)·낙관(Bull) 세 가지 가정으로 가치를 각각 산출해 결과의 범위를 제시하는 기법.",
     "날씨예보처럼 '최악·보통·최상'을 모두 보여줘 한 가지 숫자에 속지 않게 하는 것.",
     "적정주가를 단일값이 아닌 '범위'로 제시하는 근거가 된다.")
term("민감도분석 / 토네이도 차트", "Sensitivity Analysis / Tornado Chart",
     "WACC·성장률 등 핵심 변수를 조금씩 바꿀 때 가치가 얼마나 변하는지 측정·시각화한 것.",
     "요리에서 소금·설탕 양을 바꿔보며 '무엇이 맛을 가장 좌우하나'를 찾는 실험.",
     "밸류에이션 대시보드에서 어떤 가정이 결과에 가장 민감한지 보여준다.")

section("1-3. 자본비용 (WACC와 그 구성요소)")
term("WACC (가중평균자본비용)", "Weighted Average Cost of Capital",
     "WACC = (E/V)·Re + (D/V)·Rd·(1−Tc). 자기자본비용(Re)과 세후 타인자본비용(Rd)을 자본구조 비중으로 가중평균한 할인율.",
     "회사가 돈을 끌어오는 '평균 이자율'. 주주 몫과 빚 몫의 비용을 섞은 종합 자본 단가.",
     "DCF의 할인율로 직접 사용. 엔진이 종목마다 자동 계산한다.")
term("CAPM (자본자산가격결정모형)", "Capital Asset Pricing Model",
     "Re = Rf + β·(Rm − Rf). 무위험수익률에 시장위험을 베타만큼 얹어 자기자본비용을 구하는 모형.",
     "'안전한 이자 + 이 주식이 시장보다 얼마나 출렁이나'를 더해 주주가 요구할 수익률을 계산.",
     "WACC의 자기자본비용(Re) 산출에 사용된다.")
term("무위험수익률", "Risk-free Rate (Rf)",
     "원금 손실 위험이 사실상 없는 자산의 수익률. 보통 국고채(장기 국채) 금리를 사용한다.",
     "은행 예금처럼 '떼일 걱정 없는' 기본 이자율. 모든 수익률 계산의 출발선.",
     "ECOS(한국은행)에서 국고채 금리를 수집해 CAPM에 입력한다.")
term("시장위험프리미엄", "Equity Risk Premium (ERP)",
     "주식시장 전체가 무위험자산보다 추가로 요구하는 평균 초과수익률(Rm − Rf).",
     "예금 대신 위험한 주식에 투자한 대가로 '더 받아야 하는 보너스 수익률'.",
     "CAPM에서 베타에 곱해지는 항으로 자기자본비용을 끌어올린다.")
term("베타", "Beta (β)",
     "개별 주식 수익률이 시장 수익률 변동에 반응하는 민감도. β=1이면 시장과 동행, >1이면 더 출렁, <1이면 덜 출렁.",
     "시장이 파도라면 베타는 '이 배가 파도에 얼마나 흔들리나'. 높을수록 멀미가 심하다.",
     "피어 그룹의 베타를 추정·보정(Blume·Hamada)해 CAPM에 사용한다.")
term("베타 보정 (Blume·Hamada)", "Beta Adjustment",
     "Blume 보정은 베타를 장기평균 1로 회귀시키고, Hamada 보정은 자본구조(부채비율) 차이를 반영해 베타를 조정한다.",
     "남의 회사 흔들림을 빌려와 우리 회사에 맞게 '체격·빚 차이를 감안해 보정'하는 것.",
     "피어 기반 베타의 신뢰도를 높여 WACC 정확도를 개선한다.")
term("자기자본비용 / 타인자본비용", "Cost of Equity / Cost of Debt",
     "자기자본비용(Re)은 주주가 요구하는 수익률(CAPM), 타인자본비용(Rd)은 차입금에 대한 실효 이자율(세금 절감 반영).",
     "주주에게 줘야 할 '기대 보상'과 은행에 줘야 할 '이자'. 둘 다 회사 입장에선 비용.",
     "두 값을 자본구조 비중으로 섞어 WACC를 만든다.")

section("1-4. 상대가치 멀티플")
term("PER (주가수익비율)", "Price-to-Earnings Ratio",
     "PER = 주가 ÷ 주당순이익(EPS). 이익 1원당 시장이 매기는 가격. 낮을수록 이익 대비 저렴.",
     "'이 회사 1년 이익의 몇 배 가격에 팔리나'. 10배면 10년 벌어야 본전.",
     "상대가치 멀티플 비교와 PER 역산 적정주가 산출에 사용.")
term("PBR (주가순자산비율)", "Price-to-Book Ratio",
     "PBR = 주가 ÷ 주당순자산(BPS). 장부상 순자산 대비 주가 수준. 1배 미만이면 청산가치보다 싸다는 신호.",
     "'회사가 가진 순재산의 몇 배에 거래되나'. 1배면 재산만큼, 0.5배면 반값.",
     "PBR 역산으로 자산가치 기반 적정주가를 보조 산출한다.")
term("PSR (주가매출비율)", "Price-to-Sales Ratio",
     "PSR = 시가총액 ÷ 매출액. 이익이 아직 적자거나 변동성이 큰 기업의 가치 비교에 유용.",
     "이익이 들쭉날쭉한 회사를 '매출 규모'로 비교하는 자.",
     "이익 기반 멀티플이 왜곡될 때 보조 지표로 활용.")
term("EV/EBITDA", "Enterprise Value to EBITDA",
     "EV(기업가치) ÷ EBITDA. 자본구조·감가상각 정책 차이를 중화해 기업 간 수익력을 비교하는 글로벌 표준 멀티플.",
     "빚·세금·감가상각 같은 '회계 화장'을 지우고 순수 영업 체력만으로 비교하는 자.",
     "4-Way 내재가치의 한 축으로 EV/EBITDA 역산 가치를 계산한다.")
term("EV (기업가치)", "Enterprise Value",
     "EV = 시가총액 + 순부채(총차입금 − 현금). 회사를 통째로 인수할 때 실제 지불해야 하는 가치.",
     "회사를 통째로 살 때 '주식값 + 떠안을 빚 − 가진 현금'. 인수 실부담액.",
     "EV/EBITDA 멀티플과 주주가치 환산의 기준이 된다.")
term("EBIT / EBITDA", "Earnings Before Interest & Taxes (and D&A)",
     "EBIT은 이자·세금 차감 전 영업이익, EBITDA는 거기에 감가상각비(D&A)까지 더해 현금창출력에 근접시킨 지표.",
     "EBIT은 '본업으로 번 돈', EBITDA는 거기에 '장부상 비용(감가상각)을 되돌려' 실제 현금에 가깝게 본 것.",
     "FCFF 계산과 EV/EBITDA 멀티플의 핵심 입력값.")
term("EPS / BPS", "Earnings/Book-value Per Share",
     "EPS = 당기순이익 ÷ 주식수(주당순이익), BPS = 순자산 ÷ 주식수(주당순자산).",
     "회사 이익·재산을 '주식 한 장당'으로 쪼갠 값. 한 주가 품은 이익과 재산.",
     "PER·PBR 역산으로 적정주가를 구할 때 곱해지는 기준값.")
term("멀티플 / 피어 그룹", "Multiple / Peer Group",
     "멀티플은 가치를 이익·매출 등으로 나눈 배수, 피어 그룹은 비교 대상이 되는 동종·유사 기업 집합.",
     "멀티플은 '몇 배냐'의 배수, 피어는 '같은 반 친구들'. 친구들 평균과 비교한다.",
     "임베딩+산업분류로 피어를 자동 선정해 상대가치를 매긴다.")

section("1-5. 재무제표와 핵심 계정")
term("재무상태표 (B/S)", "Balance Sheet",
     "특정 시점의 자산 = 부채 + 자본을 보여주는 재무제표. 기업의 '재산 상태' 스냅샷.",
     "어느 날 통장·집·빚을 한 장에 적은 '재산 명세서'. 가진 것과 갚을 것의 균형.",
     "종목 상세에서 3개년 B/S를 DART 원본 기준으로 제공한다.")
term("손익계산서 (I/S)", "Income Statement",
     "일정 기간의 매출·비용·이익 흐름을 보여주는 재무제표. 매출액→영업이익→당기순이익으로 내려간다.",
     "한 해 동안 '얼마 벌어 얼마 쓰고 얼마 남겼나'를 적은 가계부.",
     "재무 하이라이트(YoY)와 손익흐름도(Sankey)의 원천 데이터.")
term("현금흐름표", "Cash Flow Statement",
     "영업·투자·재무 활동별 현금의 유입·유출을 보여주는 재무제표. 이익과 실제 현금의 차이를 드러낸다.",
     "'장부상 이익'이 아니라 '진짜 통장에 들어온 현금'의 흐름을 추적한 표.",
     "FCFF 산출과 재무 건전성 판단의 근거.")
term("감가상각비 / 자본적지출", "D&A / CapEx",
     "감가상각비(D&A)는 설비 가치를 기간에 걸쳐 비용 처리한 것(현금 유출 없음), CapEx는 설비·자산에 실제 투자한 현금 지출.",
     "감가상각은 '장비가 닳은 만큼 장부에서 깎는 것', CapEx는 '새 장비 사는 데 쓴 진짜 돈'.",
     "FCFF 공식에서 D&A는 더하고 CapEx는 빼서 실제 현금을 구한다.")
term("순운전자본", "Net Working Capital (NWC)",
     "유동자산 − 유동부채. 영업을 돌리는 데 묶이는 단기 자금(재고·매출채권 등).",
     "장사하려고 '재고와 외상값에 깔아둔 돈'. 늘면 현금이 묶이고, 줄면 현금이 풀린다.",
     "FCFF에서 ΔNWC(증감)를 빼 실제 가용 현금을 반영한다.")
term("손익흐름도", "Income Sankey Diagram",
     "매출에서 각종 비용을 거쳐 이익에 이르는 흐름을 폭으로 시각화한 다이어그램(Sankey).",
     "물줄기가 갈라지듯 '매출이라는 큰 물이 비용으로 새고 이익으로 남는' 과정을 한눈에.",
     "약 2,627종목의 손익흐름도를 생성해 기업개요 탭에 제공한다.")

section("1-6. 재무비율 · 건전성 · 시장 지표")
term("ROE / ROA / ROIC", "Return on Equity/Assets/Invested Capital",
     "투입 자본 대비 수익성 지표. ROE=순이익/자기자본, ROA=순이익/총자산, ROIC=세후영업이익/투하자본.",
     "'넣은 돈 대비 얼마나 잘 버나'의 성적표. 높을수록 자본을 효율적으로 굴린다.",
     "가치진단과 재무 건전성 신호등의 판단 근거로 사용.")
term("부채비율", "Debt-to-Equity (D/E)",
     "부채비율 = 총부채 ÷ 자기자본 × 100(%). 재무 안정성·레버리지 수준을 나타낸다.",
     "'내 돈 대비 빚이 얼마나 많나'. 높으면 위험, 낮으면 안정적이지만 성장은 느릴 수 있다.",
     "재무 건전성 신호등(위험 경고)과 Hamada 베타 보정에 활용.")
term("이자보상배율", "Interest Coverage Ratio",
     "영업이익 ÷ 이자비용. 영업으로 번 돈으로 이자를 몇 배 감당할 수 있는지. 1배 미만이면 위험.",
     "'버는 돈으로 이자를 몇 번 낼 수 있나'. 낮으면 빚 갚기 빠듯하다는 경고.",
     "재무 건전성 위험신호 판정 지표로 사용.")
term("신용등급", "Credit Rating",
     "신용평가사가 기업의 채무상환능력을 등급(AAA~D)으로 매긴 것. 타인자본비용 추정에 쓰인다.",
     "회사의 '신용점수'. 높을수록 돈을 싸게 빌리고, 낮으면 비싸게 빌린다.",
     "타인자본비용(Rd) 추정과 RAG 기반 신용정보 추출에 활용.")
term("배당", "Dividend",
     "기업이 이익의 일부를 주주에게 현금·주식으로 분배하는 것. 배당수익률 = 주당배당금 ÷ 주가.",
     "회사가 번 돈을 주주에게 '용돈처럼 나눠주는' 것.",
     "기업개요·재무 정보의 일부로 제공한다.")
term("컨센서스", "Consensus Estimate",
     "다수 증권사 애널리스트 추정치의 평균. 시장의 기대를 집약한 기준선.",
     "여러 전문가의 예측을 '평균 낸 다수 의견'. 실적이 이보다 좋으면 어닝 서프라이즈.",
     "애널리스트 리포트 수집·반영으로 추정의 현실성을 보강한다.")
term("YoY / QoQ", "Year-over-Year / Quarter-over-Quarter",
     "YoY는 전년 동기 대비, QoQ는 전분기 대비 증감률. 계절성을 통제하려면 YoY가 유용.",
     "'작년 같은 때보다(YoY)' 혹은 '지난 분기보다(QoQ)' 얼마나 변했나.",
     "재무 하이라이트에 YoY 라벨로 증감을 표시한다.")
term("보통주 / 우선주", "Common / Preferred Stock",
     "보통주는 의결권이 있는 일반 주식, 우선주는 의결권이 없는 대신 배당이 우선되는 주식.",
     "보통주는 '투표권 있는 회원', 우선주는 '투표권 없지만 배당 먼저 받는 회원'.",
     "검색은 본주(보통주) 중심으로 노출해 혼선을 줄인다.")
term("발행주식수 / 시가총액", "Shares Outstanding / Market Cap",
     "발행주식수는 시장에 풀린 총 주식 수, 시가총액 = 주가 × 발행주식수(회사 전체의 시장 가격).",
     "시가총액은 '주식을 전부 사들이는 데 드는 돈'. 회사의 몸값.",
     "EPS·BPS·EV 계산과 종목 정렬(대장주 우선)의 기준.")
term("가치진단 (고평가/저평가)", "Valuation Diagnosis",
     "적정가치 대비 현재가의 위치로 고평가·적정·저평가를 판정하는 종합 진단.",
     "'지금 비싸게 사는 건지 싸게 사는 건지'를 신호등처럼 알려주는 것.",
     "LLM 해설과 함께 스코어카드로 직관적으로 제시한다.")

doc.add_page_break()
# ════════════════════════════════════════════════════════
# PART 2 — 개발·기술·인프라
# ════════════════════════════════════════════════════════
H("PART 2.  개발 · 기술 · 인프라 용어", 18, NAVY, 4, 8)

section("2-1. 생성형 AI와 RAG")
term("LLM (대규모 언어모델)", "Large Language Model",
     "방대한 텍스트로 학습해 사람처럼 문장을 이해·생성하는 인공지능 모델(예: GPT 계열).",
     "엄청난 책을 다 읽고 말을 배운 'AI 두뇌'. 질문하면 문장으로 답한다.",
     "챗봇 답변 생성과 히스토리 브리핑 요약에 사용(gpt-5-mini, 보조 gpt-4o-mini).")
term("토큰 / 컨텍스트 윈도우", "Token / Context Window",
     "토큰은 모델이 처리하는 텍스트의 최소 단위(대략 단어 조각), 컨텍스트 윈도우는 한 번에 다룰 수 있는 토큰의 최대량.",
     "토큰은 '글자 조각', 컨텍스트 윈도우는 'AI가 한 번에 기억할 수 있는 분량'.",
     "긴 사업보고서를 청크로 나눠 윈도우 한계 안에서 처리한다.")
term("임베딩 / 벡터", "Embedding / Vector",
     "텍스트의 의미를 수백 개 숫자의 배열(벡터)로 바꾼 것. 의미가 가까우면 벡터도 가깝다.",
     "글의 '의미를 좌표로 찍은 것'. 비슷한 뜻이면 지도에서 가까운 위치에 놓인다.",
     "사업보고서를 임베딩해 챗봇이 의미 기반으로 검색하게 한다.")
term("벡터 데이터베이스", "Vector Database",
     "임베딩 벡터를 저장하고 '가장 비슷한 것'을 빠르게 찾아주는 특수 데이터베이스(예: ChromaDB).",
     "뜻이 비슷한 문서를 순식간에 찾아주는 'AI 전용 도서관'.",
     "약 40GB 규모의 ChromaDB에 사업보고서 청크(약 192만 개)를 저장한다.")
term("코사인 유사도", "Cosine Similarity",
     "두 벡터가 이루는 각도로 의미적 유사성을 0~1로 측정하는 방법. 1에 가까울수록 유사.",
     "두 화살표가 같은 방향이면 '비슷한 뜻', 다른 방향이면 '다른 뜻'으로 보는 자.",
     "질문과 가장 가까운 문서 청크를 골라내는 검색 기준.")
term("청크", "Chunk",
     "긴 문서를 검색·임베딩하기 좋게 일정 크기로 잘게 나눈 조각.",
     "두꺼운 책을 '페이지·문단 단위로 잘라' 필요한 부분만 빨리 찾게 한 것.",
     "사업보고서를 청크로 분할해 벡터DB에 적재한다.")
term("RAG (검색증강생성)", "Retrieval-Augmented Generation",
     "LLM이 답하기 전에 신뢰할 수 있는 문서를 먼저 검색해, 그 근거 위에서 답을 생성하는 구조. 환각을 줄인다.",
     "AI가 '아무 말'을 하지 않도록 '관련 자료를 먼저 찾아보고' 답하게 시키는 방식.",
     "AI 챗봇의 핵심 아키텍처. 답변마다 출처 카드를 함께 제시한다.")
term("리랭킹", "Reranking",
     "1차 검색으로 추린 후보 문서를 더 정교한 모델로 재정렬해 가장 적합한 근거를 위로 올리는 단계.",
     "검색 결과를 '한 번 더 꼼꼼히 줄 세워' 진짜 관련 있는 걸 맨 위로 보내는 것.",
     "GPU 리랭커로 답변 정확도를 높인다(로컬 연산, 무료).")
term("프롬프트", "Prompt",
     "LLM에게 주는 입력(지시문·맥락·질문). 프롬프트 설계가 답변 품질을 좌우한다.",
     "AI에게 시키는 '주문서'. 잘 적을수록 원하는 답이 나온다.",
     "검색된 근거+질문을 결합한 프롬프트로 답을 생성한다.")
term("할루시네이션", "Hallucination",
     "LLM이 근거 없이 그럴듯한 거짓 정보를 생성하는 현상.",
     "AI가 모르는 걸 '아는 척 지어내는' 것. 가장 경계해야 할 오류.",
     "NO-MOCK 원칙·RAG 출처표시·검증 페이지로 원천 차단한다.")
term("파인튜닝 vs RAG", "Fine-tuning vs RAG",
     "파인튜닝은 모델 자체를 추가 학습시키는 것, RAG는 모델은 그대로 두고 외부 지식을 검색해 주입하는 것.",
     "파인튜닝은 'AI를 다시 교육', RAG는 'AI에게 참고서를 쥐여주기'. 후자가 빠르고 갱신이 쉽다.",
     "최신 공시를 반영해야 하므로 RAG 방식을 채택했다.")

section("2-2. 백엔드 · API")
term("API / REST API", "Application Programming Interface",
     "프로그램끼리 약속된 형식으로 데이터를 주고받는 창구. REST는 HTTP 기반의 대표적 설계 양식.",
     "식당의 '주문 창구'. 정해진 메뉴(요청)를 넣으면 음식(데이터)이 나온다.",
     "프론트엔드가 백엔드 API를 호출해 모든 데이터를 받아온다.")
term("엔드포인트", "Endpoint",
     "API에서 특정 기능에 접근하는 고유 주소(URL 경로). 예: /api/overview/{종목코드}.",
     "주문 창구 중 '이 메뉴 전용 줄'. 주소마다 받는 데이터가 다르다.",
     "종목별 개요·재무·밸류·공시 등 기능마다 엔드포인트가 있다.")
term("FastAPI", "FastAPI",
     "파이썬 기반의 고성능 비동기 웹 API 프레임워크. 자동 문서화·타입 검증을 지원한다.",
     "백엔드를 빠르고 튼튼하게 짓게 해주는 '파이썬 건축 키트'.",
     "FINSIGHT 백엔드 서버(8090 포트)가 FastAPI로 구현돼 있다.")
term("JSON", "JavaScript Object Notation",
     "키-값 쌍으로 데이터를 표현하는 경량 텍스트 포맷. API 통신의 사실상 표준.",
     "'이름: 값' 형태로 데이터를 적은 만국공통 메모 양식.",
     "API 응답·밸류에이션 결과 저장 등 데이터 교환 전반에 사용.")
term("HTTP 상태코드", "HTTP Status Code",
     "요청 처리 결과를 숫자로 알리는 표준 코드. 200(성공)·404(없음)·500(서버오류) 등.",
     "택배 송장의 '배송완료·주소없음·사고' 같은 상태 표시.",
     "전수 검수에서 500(버그)·404(데이터없음)를 분류해 품질을 관리한다.")
term("CORS", "Cross-Origin Resource Sharing",
     "다른 도메인 간 브라우저 요청을 허용·제한하는 보안 정책.",
     "'우리 가게 음식을 다른 동네에 배달해도 되나'를 정하는 규칙.",
     "프론트(3000)와 챗봇(8800) 간 통신을 위해 CORS를 설정했다.")
term("비동기 처리", "Asynchronous",
     "한 작업의 완료를 기다리지 않고 다른 작업을 동시에 진행하는 방식. 응답성을 높인다.",
     "라면 끓이며 동시에 설거지하는 것. 하나 끝날 때까지 멈춰 서지 않는다.",
     "여러 API를 병렬 호출해 종목 상세 화면 로딩을 빠르게 한다.")

section("2-3. 프론트엔드")
term("Next.js / React", "Next.js / React",
     "React는 UI를 컴포넌트로 만드는 자바스크립트 라이브러리, Next.js는 React에 서버렌더링·라우팅을 더한 프레임워크.",
     "화면을 '레고 블록(컴포넌트)'으로 조립하게 해주는 도구와 그 확장판.",
     "FINSIGHT 프론트엔드 전체가 Next.js 16(React 19)로 구현돼 있다.")
term("컴포넌트 / 상태", "Component / State",
     "컴포넌트는 재사용 가능한 UI 조각, 상태(state)는 그 조각이 기억하는 변하는 데이터.",
     "컴포넌트는 '레고 블록', 상태는 그 블록이 든 '현재 값(예: 켜짐/꺼짐)'.",
     "밸류 탭·차트·토글 등이 모두 컴포넌트로 구성된다.")
term("SSR / CSR", "Server/Client-Side Rendering",
     "SSR은 서버가 완성된 화면을 보내 첫 로딩과 SEO에 유리, CSR은 브라우저가 화면을 그려 상호작용에 유리.",
     "SSR은 '완성된 요리를 배달', CSR은 '재료를 받아 집에서 조리'. 상황에 맞게 섞어 쓴다.",
     "초기 화면은 빠르게, 상호작용은 매끄럽게 하도록 혼합 적용.")
term("스켈레톤 UI", "Skeleton UI",
     "데이터 로딩 중 실제 콘텐츠 자리에 회색 뼈대를 먼저 보여줘 체감 속도를 높이는 기법.",
     "음식 나오기 전 '자리에 깔아두는 빈 접시'. 기다림이 덜 답답하게 느껴진다.",
     "글로벌 마켓 시그널 위젯 로딩에 스켈레톤을 적용했다.")
term("반응형 / 다크모드", "Responsive / Dark Mode",
     "반응형은 화면 크기에 맞춰 레이아웃이 자동 조정되는 설계, 다크모드는 어두운 배색 테마.",
     "옷이 몸에 맞게 늘어나는 것(반응형)과 야간용 검은 화면(다크모드).",
     "PC·모바일 대응과 사용자 선호 테마를 지원한다.")

section("2-4. 데이터베이스 · 데이터 모델")
term("SQLite / PostgreSQL", "SQLite / PostgreSQL",
     "둘 다 관계형 DB(SQL). SQLite는 파일 하나로 가벼운 로컬용, PostgreSQL은 대규모·동시접속에 강한 서버용.",
     "SQLite는 '노트 한 권', PostgreSQL은 '도서관'. 규모가 커지면 도서관으로 옮긴다.",
     "로컬 SQLite로 개발 후 AWS RDS PostgreSQL로 이관(16테이블·약 1,175만 행).")
term("RDS", "Relational Database Service",
     "AWS가 관리해주는 관계형 데이터베이스 서비스. 백업·확장·보안을 자동화한다.",
     "직접 관리 안 해도 되는 '클라우드가 운영해주는 도서관'.",
     "기업개요·재무·밸류 정형 데이터를 RDS PostgreSQL에 저장한다.")
term("MongoDB", "MongoDB",
     "표(테이블) 대신 유연한 문서(JSON 유사) 단위로 저장하는 NoSQL 데이터베이스.",
     "칸이 고정된 표가 아니라 '자유로운 양식의 서류함'. 형태가 다양한 데이터에 유리.",
     "AI 히스토리 브리핑(서술형)을 MongoDB Atlas에 저장한다.")
term("ChromaDB", "ChromaDB",
     "임베딩 벡터를 저장·검색하는 오픈소스 벡터 데이터베이스.",
     "뜻이 비슷한 글을 찾아주는 'AI 전용 무료 도서관'.",
     "RAG 챗봇의 검색 엔진. 약 40GB 규모로 EC2(GPU)에서 운영한다.")
term("S3", "Simple Storage Service",
     "AWS의 대용량 객체(파일) 저장소. 이미지·HTML·JSON 등을 안정적으로 보관·서빙한다.",
     "무한히 큰 '클라우드 창고'. 사진·파일을 넣어두고 링크로 꺼내 쓴다.",
     "Sankey HTML·밸류 JSON·계열사 이미지를 S3에 보관·서빙한다.")
term("폴리글랏 퍼시스턴스", "Polyglot Persistence",
     "데이터 성격에 따라 서로 다른 종류의 DB를 함께 쓰는 설계 전략.",
     "'연장통'에서 못엔 망치, 나사엔 드라이버를 쓰듯 데이터마다 맞는 DB를 쓰는 것.",
     "정형(RDS)·문서(Mongo)·벡터(Chroma)·파일(S3) 4종을 목적별로 분리했다.")
term("스키마 / 테이블 / PK·FK", "Schema / Table / Primary·Foreign Key",
     "스키마는 DB의 구조 설계, 테이블은 행·열로 된 데이터 묶음, PK는 행을 식별하는 고유키, FK는 다른 테이블을 참조하는 키.",
     "스키마는 '설계도', 테이블은 '엑셀 시트', PK는 '학번', FK는 '다른 명단을 가리키는 연결고리'.",
     "종목코드(PK)로 여러 테이블을 연결해 한 종목 정보를 모은다.")
term("ERD", "Entity-Relationship Diagram",
     "데이터베이스의 테이블(개체)과 그들 간 관계를 도식화한 설계도.",
     "회사 조직도처럼 '데이터들이 어떻게 연결되는지' 그린 지도.",
     "PM 산출물로 4-DB 구조의 ERD를 제공한다(원형 아님, 정석 테이블-관계).")
term("정규화 / 인덱스", "Normalization / Index",
     "정규화는 중복을 줄이도록 테이블을 분리·정리하는 것, 인덱스는 검색 속도를 높이는 색인 구조.",
     "정규화는 '같은 정보를 여러 곳에 안 적기', 인덱스는 책 뒤의 '찾아보기'.",
     "재무·종목 테이블을 정규화하고 종목코드에 인덱스를 둬 빠르게 조회한다.")
term("마이그레이션", "Migration",
     "데이터·스키마를 한 환경(DB)에서 다른 환경으로 옮기는 작업.",
     "이삿짐을 '하나도 빠뜨리지 않고' 새집으로 옮기는 것.",
     "로컬 SQLite → AWS RDS로 이관하며 행 수 일치를 검증했다(NO-MOCK).")

section("2-5. 클라우드 · 인프라 · 운영")
term("AWS / 클라우드", "Amazon Web Services / Cloud",
     "인터넷으로 서버·저장소·DB 등 IT 자원을 빌려 쓰는 서비스. 직접 장비를 사지 않아도 된다.",
     "건물을 짓지 않고 '필요한 만큼 사무실을 임대'하는 것. 쓴 만큼만 낸다.",
     "FINSIGHT 전체를 AWS에 배포해 외부에서 접속 가능하게 했다.")
term("EC2", "Elastic Compute Cloud",
     "AWS의 가상 서버. 운영체제·앱을 올려 실제 프로그램을 구동한다.",
     "클라우드에서 빌리는 '컴퓨터 한 대'. 켜면 과금, 끄면 멈춘다.",
     "백엔드·프론트·GPU 챗봇을 EC2에서 운영한다.")
term("GPU 인스턴스", "GPU Instance",
     "그래픽처리장치(GPU)를 갖춰 AI 연산을 빠르게 처리하는 고성능 서버.",
     "AI 계산을 위한 '근육질 컴퓨터'. 비싸므로 필요할 때만 켠다.",
     "챗봇 임베딩·리랭킹용. 데모 시에만 켜고 끝나면 정지해 비용을 아낀다.")
term("EBS / Elastic IP", "Elastic Block Store / Elastic IP",
     "EBS는 EC2에 붙이는 영구 디스크(인스턴스를 꺼도 데이터 보존), Elastic IP는 고정 공인 IP 주소.",
     "EBS는 '꺼도 안 지워지는 외장하드', Elastic IP는 '바뀌지 않는 우리 집 주소'.",
     "ChromaDB를 EBS에 둬 인스턴스를 정지해도 데이터가 유지된다.")
term("nginx / 리버스 프록시", "nginx / Reverse Proxy",
     "nginx는 고성능 웹서버, 리버스 프록시는 외부 요청을 받아 내부 서버로 분배·중계하는 역할.",
     "건물 1층의 '안내데스크'. 손님 요청을 받아 알맞은 부서로 연결한다.",
     "/api는 백엔드, /는 프론트로 분배하고 HTTPS를 처리한다.")
term("HTTPS / SSL·TLS / 인증서", "HTTPS / SSL·TLS / Certificate",
     "HTTPS는 암호화된 안전한 웹 통신, SSL·TLS는 그 암호화 프로토콜, 인증서는 신원을 보증하는 전자 증명서.",
     "택배를 '자물쇠 달린 상자'로 보내고, 보낸 이가 진짜임을 보증하는 '인감'을 붙이는 것.",
     "Let's Encrypt 인증서로 배포 사이트를 https로 제공한다.")
term("systemd / 스케줄러(cron)", "systemd / Scheduler",
     "systemd는 서버 부팅 시 서비스를 자동 실행·관리하는 도구, 스케줄러(cron)는 정해진 시각에 작업을 자동 실행한다.",
     "systemd는 '컴퓨터 켜지면 알아서 가게 문 열기', cron은 '매일 정해진 시각의 알람'.",
     "백엔드·프론트 자동 기동, 주가 데이터를 매 거래일 16:00 자동 갱신한다.")
term("보안그룹 / 가용영역", "Security Group / Availability Zone",
     "보안그룹은 서버의 인·아웃 트래픽을 통제하는 방화벽, 가용영역은 물리적으로 분리된 데이터센터 구역.",
     "보안그룹은 '출입 통제 명단', 가용영역은 '재해에 대비한 분산 건물'.",
     "DB 접근을 화이트리스트로 제한하고 일배치 시 자기 IP만 일시 허용한다.")
term("헬스체크 / 멱등성", "Health Check / Idempotency",
     "헬스체크는 서비스 정상 동작을 주기적으로 확인하는 것, 멱등성은 같은 작업을 여러 번 해도 결과가 같은 성질.",
     "헬스체크는 '맥박 재기', 멱등성은 '엘리베이터 버튼을 여러 번 눌러도 한 번과 같은' 안전성.",
     "전수 검수로 전 종목 헬스체크를 수행하고, 데이터 적재는 멱등하게 설계했다.")

section("2-6. 데이터 출처 · 포맷 · 협업")
term("DART / OpenAPI", "DART / Open API",
     "DART는 금융감독원 전자공시시스템, OpenAPI는 외부에 공개된 데이터 제공 인터페이스.",
     "DART는 기업 공시의 '공식 도서관', OpenAPI는 그 자료를 자동으로 받아오는 '창구'.",
     "재무·사업보고서·공시를 DART OpenAPI로 수집한다.")
term("XBRL", "eXtensible Business Reporting Language",
     "재무 데이터를 기계가 읽을 수 있도록 표준 태그로 구조화한 국제 전자공시 포맷.",
     "재무제표에 '바코드를 붙여' 컴퓨터가 자동으로 읽게 만든 것.",
     "사업보고서 XBRL을 파싱해 3개년 재무를 정확히 적재한다.")
term("OHLCV", "Open-High-Low-Close-Volume",
     "특정 기간의 시가·고가·저가·종가·거래량 묶음. 주가 차트의 기본 데이터.",
     "하루 주가의 '시작값·최고·최저·끝값·거래량' 다섯 숫자.",
     "일봉 OHLCV를 매 거래일 자동 동기화해 차트로 보여준다.")
term("yfinance / ECOS / WICS", "Data Sources",
     "yfinance는 글로벌 시세 라이브러리, ECOS는 한국은행 경제통계(금리), WICS는 산업분류 체계(FnGuide).",
     "각각 '실시간 시세 창구', '금리 통계 창구', '업종 분류표'.",
     "실시간 시세·국고채금리·산업분류 데이터를 이들에서 받는다.")
term("CSV", "Comma-Separated Values",
     "쉼표로 칸을 나눈 단순 표 형식의 텍스트 파일. 데이터 교환의 기본 포맷.",
     "엑셀을 '메모장으로 저장한' 가장 단순한 표.",
     "수집 단계의 중간 저장·종목 목록 등에 사용한다.")
term("Git / GitHub", "Git / GitHub",
     "Git은 코드 변경 이력을 관리하는 버전관리 시스템, GitHub은 이를 온라인에서 협업하는 플랫폼.",
     "Git은 글의 '수정 기록·되돌리기', GitHub은 여러 명이 함께 쓰는 '공유 작업실'.",
     "팀 코드 통합·이력 관리에 사용하며, 비밀키는 절대 올리지 않는다.")
term("CI/CD", "Continuous Integration/Delivery",
     "코드 변경을 자동으로 통합·검증·배포하는 개발 자동화 파이프라인.",
     "코드를 고치면 '자동으로 검사하고 배포까지' 흘러가는 컨베이어벨트.",
     "빌드·배포 절차를 표준화해 배포 실수를 줄인다.")
term("MCP", "Model Context Protocol",
     "외부 AI(LLM)가 우리 서비스의 기능을 표준화된 '도구'로 호출하게 해주는 개방형 프로토콜.",
     "AI가 우리 서비스를 '플러그인처럼 꽂아 쓰게' 하는 만국공통 콘센트.",
     "FINSIGHT MCP 서버로 종목검색·밸류·기업개요·브리핑을 도구로 노출한다.")
term("캐시", "Cache",
     "자주 쓰는 데이터를 빠르게 꺼내도록 임시 저장해두는 것.",
     "자주 보는 책을 '책상 위에 미리 꺼내두는' 것. 매번 서고에 갈 필요가 없다.",
     "마켓 시그널·반복 조회 결과를 캐시해 응답을 빠르게 한다.")

doc.add_paragraph()
body("— 본 용어집은 FINSIGHT 서비스 이해를 돕기 위한 교육용 자료입니다. 정의는 일반적 통용 기준을 따랐으며, 실제 투자 판단의 책임은 이용자에게 있습니다. —",
     8.5, GRAY, True)

out = Path(r"C:\Users\Admin\FILMN9\통합산출물\FINSIGHT_용어사전_백과사전.docx")
doc.save(out)
# 용어 수 집계
import re
txt = "\n".join(p.text for p in doc.paragraphs)
n = txt.count("▸")
print(f"저장: {out}")
print(f"총 용어 수: {n}개")
