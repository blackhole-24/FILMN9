# -*- coding: utf-8 -*-
"""AWS 배포 가이드 → Word(.docx) 생성."""
import sys
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx 미설치 →  pip install python-docx")

OUT = Path(r"C:\Users\Admin\FILMN9\통합산출물\AWS배포\00_AWS_배포_로드맵_초보자가이드.docx")
doc = Document()

# 기본 폰트
st = doc.styles["Normal"]
st.font.name = "맑은 고딕"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

def shade(p, color="F2F4F7"):
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), color)
    pPr.append(sh)

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.name = "Consolas"; r.font.size = Pt(9.5)
    r.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    p.paragraph_format.left_indent = Inches(0.3); shade(p)
    return p

def chk(text, bold_head=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run("☐  " + text);
    if bold_head: r.bold = True
    return p

def note(text):
    p = doc.add_paragraph(); r = p.add_run("💡 " + text); r.italic = True
    r.font.color.rgb = RGBColor(0x47, 0x55, 0x69); shade(p, "FFF7E6")

# 제목
t = doc.add_heading("FINSIGHT AWS 배포 로드맵 — 초보자 가이드", level=0)
sub = doc.add_paragraph("2026-06-14 · 하이브리드 구조: RDS(관계형) + EC2-GPU(ChromaDB) + Atlas(MongoDB) + S3(파일)")
sub.runs[0].italic = True; sub.runs[0].font.color.rgb = RGBColor(0x64,0x74,0x8b)
note("프론트 UI(밸류탭·챗봇·메인)는 나중에 바뀌어도 됨 — 인프라 먼저 깔고, 프론트는 마지막에 다시 빌드해서 얹으면 됨(재배포 자유). 리전은 전부 서울(ap-northeast-2).")

doc.add_heading("사전 준비 (내 PC에 설치)", level=1)
chk("AWS CLI 설치 — aws.amazon.com/cli  → 설치 후 'aws configure'로 키 입력")
chk("PostgreSQL 클라이언트(psql·pg_dump) 설치 — 설치 시 'Command Line Tools' 체크 (버전 17 권장)")
chk("pip install boto3  (S3 업로드 스크립트용, FINSIGHT_env)")
chk("AWS 콘솔 로그인/가입 — 결제수단 등록 필요")

doc.add_heading("Phase 0 · 기반 (0.5일)", level=1)
chk("IAM 사용자 생성(관리자 권한) → 액세스 키 발급 → 'aws configure' 입력")
chk("VPC: 기본 VPC 사용(처음엔 단순하게)")
chk("보안그룹 2개: finsight-app-sg(80/443/22) · finsight-db-sg(5432, app-sg에서만)")

doc.add_heading("Phase 1 · RDS (관계형 DB) (0.5~1일)", level=1)
chk("RDS → PostgreSQL 생성 (엔진 17, 서울, db.t3.medium, 50GB gp3, 보안그룹 finsight-db-sg)")
chk("현재 Supabase에서 데이터 덤프 (약 2.7GB):")
code('pg_dump "현재_DATABASE_URL(.env값)" -Fc -f finsight.dump')
chk("RDS로 복원:")
code("pg_restore -h <RDS엔드포인트> -U postgres -d postgres --no-owner finsight.dump")
chk("검증: psql ... \"SELECT count(*) FROM ohlcv;\" → 10,678,102 확인")
chk("백엔드 .env 교체: DB_HOST/PORT/NAME/USER/PASS/DATABASE_URL을 RDS 값으로 (DB_BACKEND=postgres 유지) → 코드 수정 0")
note("빠른 길: RDS 안 쓰고 Supabase 그대로 유지하면 Phase 1 전체 생략 → 배포 1~2일 단축.")

doc.add_heading("Phase 2 · S3 (정적 파일) (0.5일)", level=1)
chk("S3 버킷 생성: finsight-static-<임의> (서울)")
chk("업로드 (이 폴더의 s3_upload_static.py):")
code("set FINSIGHT_S3_BUCKET=finsight-static-xxxx\npython s3_upload_static.py")
chk("→ Sankey(2,691) + 밸류 결과(13,344) 업로드")
chk("백엔드 파일 서빙 경로를 S3(또는 CloudFront) URL로 전환")

doc.add_heading("Phase 3 · EC2-GPU + ChromaDB (1~1.5일, 가장 무거움) ⚠️", level=1)
chk("EC2 GPU 인스턴스(g4dn.xlarge 등, Ubuntu, app-sg) 생성")
chk("EBS 볼륨 60GB+ 부착 (ChromaDB 40GB용)")
chk("로컬 ChromaDB(chatbot/embedding/chroma_db ~40GB)를 EC2로 전송 (scp 또는 S3 경유)")
chk("GPU 드라이버·CUDA·임베딩/리랭커 세팅, 챗봇 서버(:8800) 기동 테스트")

doc.add_heading("Phase 4 · EC2 앱 배포 (0.5~1일)", level=1)
chk("EC2(앱): FastAPI(uvicorn) + nginx 리버스프록시")
chk("프론트: npm run build → EC2 서빙 (또는 S3+CloudFront)")
chk("MongoDB Atlas: Network Access에 EC2 공인 IP 화이트리스트 추가 (코드 변경 0)")
chk("환경변수 13개를 EC2 .env 또는 Secrets Manager에 등록")

doc.add_heading("Phase 5 · 도메인·HTTPS·검증 (0.5일)", level=1)
chk("Route53 도메인 + ACM 인증서(HTTPS) → ALB/nginx 연결")
chk("전수 헬스체크·회귀검증(NO-MOCK)")
chk("데모 URL 공유")

doc.add_heading("환경변수 목록 (Secrets Manager/.env — 값은 절대 코드/Git에 X)", level=1)
p = doc.add_paragraph()
p.add_run("DART_API_KEY · ECOS_API_KEY · OPENAI_API_KEY · DB_HOST · DB_PORT · DB_NAME · "
          "DB_USER · DB_PASS · DATABASE_URL · DB_BACKEND · MONGO_URI · MONGO_DB · "
          "MONGO_COLLECTION · Naver_Client_Id · Naver_Client_Secret")

doc.add_heading("요약", level=1)
chk("총 예상: 3~5일 (Supabase 유지 시 2~3일)", bold_head=True)
chk("월 비용 개략: RDS(t3.medium ~$60) + EC2-GPU(g4dn ~$380 상시/스팟 절감) + S3(~$1) + Atlas(기존). GPU가 최대 — 데모 기간만 켜고 끄면 절감.")
p = doc.add_paragraph(); r = p.add_run("다음에 클로드가 도와줄 것: pg_dump/restore 실행 보조 · S3 업로드 · 백엔드 .env 전환 · "
        "EC2 nginx 설정 · 헬스체크. (AWS 콘솔 클릭은 본인, 스크립트·검증은 클로드)")
r.italic = True

doc.save(OUT)
print(f"저장: {OUT}")
