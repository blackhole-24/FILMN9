# -*- coding: utf-8 -*-
"""관리자 페이지 시연 스크립트 → Word(docx) 생성. 창B."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(r"C:\Users\Admin\FILMN9\통합산출물\발표PPT\시연스크립트_관리자페이지_창B.docx")
NAVY = RGBColor(0x1F, 0x3A, 0x5F); BLUE = RGBColor(0x2F, 0x55, 0x97)
GRAY = RGBColor(0x64, 0x74, 0x8B); INK = RGBColor(0x0F, 0x17, 0x2A)
HDRBG = "2F5597"; ALT = "F2F6FC"

doc = Document()
# 기본 폰트
st = doc.styles["Normal"]; st.font.name = "맑은 고딕"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
for m in ("top_margin","bottom_margin","left_margin","right_margin"):
    setattr(sec, m, Inches(0.8))


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def setfont(run, size=10.5, bold=False, color=INK, name="맑은 고딕"):
    run.font.name = name; run.font.size = Pt(size); run.font.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def h(text, size=15, color=NAVY, before=14, after=6):
    p = doc.add_paragraph(); p.space_before = Pt(before); p.space_after = Pt(after)
    r = p.add_run(text); setfont(r, size, True, color); return p


def para(text, size=10.5, color=INK, italic=False, before=2, after=2):
    p = doc.add_paragraph(); p.space_before = Pt(before); p.space_after = Pt(after)
    r = p.add_run(text); setfont(r, size, False, color); r.font.italic = italic
    return p


def quote(text):
    p = doc.add_paragraph(); p.space_before = Pt(4); p.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run("🎙 " + text); setfont(r, 10.5, False, RGBColor(0x33,0x40,0x55)); r.font.italic = True
    return p


def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.width = Inches(widths[i]); shade(c, HDRBG)
        c.paragraphs[0].clear(); r = c.paragraphs[0].add_run(hd); setfont(r, 10, True, RGBColor(0xFF,0xFF,0xFF))
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].width = Inches(widths[ci])
            if ri % 2 == 1: shade(cells[ci], ALT)
            cells[ci].paragraphs[0].clear()
            r = cells[ci].paragraphs[0].add_run(val)
            setfont(r, 9.5, ci == 0, BLUE if ci == 0 else INK)
    return t

# ─── 제목 ───
tp = doc.add_paragraph(); tp.space_after = Pt(2)
r = tp.add_run("FINSIGHT 관리자 페이지 — 시연 발표 스크립트 (도슨트형)"); setfont(r, 18, True, NAVY)
para("담당: 창B(PM)  ·  화면: https://43.203.94.124.nip.io/admin  ·  4탭  ·  2026-06-19", 9, GRAY)
para("작성 근거: 실제 admin 소스(frontend/app/admin/) + 태상님 밸류에이션 운영매뉴얼", 9, GRAY)
para("톤: 미술관 도슨트처럼 — 각 박스가 무엇을(역할) · 왜 보는지(의미) · 어떻게 읽는지(해석). 비전공자도 한 문장으로 이해.", 9.5, GRAY, italic=True)
kp = doc.add_paragraph(); kp.space_before = Pt(4)
kr = kp.add_run("★ 관통 메시지: 화면만 만든 게 아니라, 그 화면이 진짜인지 스스로 감시·검증하는 장치까지 만들었습니다. (NO-MOCK)")
setfont(kr, 10.5, True, BLUE)

# ─── 0. 인트로 ───
h("0. 들어가며 (관리자 페이지란?) — 30초")
quote("지금까지 보신 건 일반 사용자 화면이었습니다. 이제 보여드릴 관리자 페이지는 저희가 이 서비스를 어떻게 운영하고, 데이터가 진짜인지 어떻게 검증하는지를 담은 내부 통제판입니다. '믿어주세요'가 아니라 '직접 확인해 보세요'라고 말씀드리는 곳입니다. 비밀번호로 보호되며, 총 4개 탭 — 기업분석 운영/검증, 밸류에이션 운영/테스트 입니다. 한 가지만 기억해 주세요. 이 서비스는 '없는 값은 지어내지 않는다(NO-MOCK)' 원칙이라, 빈칸이나 '준비 중'은 고장이 아니라 '아직 데이터가 없다'는 정직한 표시입니다.")

# ─── 탭1 ───
h("탭 1. 기업분석 · 운영 — '지금 서비스가 살아있고, 데이터가 최신인가'")
para("첫 번째 탭은 운영 상태판입니다. 서비스가 정상인지, 데이터가 며칠 자인지를 위에서 아래로 훑어봅니다.")
table(["박스", "역할 (무엇을)", "의미·해석 (어떻게 읽나)"], [
 ["🗺️ 기능 모듈 운영 현황", "메인·기업개요·밸류·챗봇 각 기능이 어느 원천→어떻게 화면까지 + 지금 정상인지", "데이터 계보(lineage). 초록=정상. 화면과 출처가 1:1로 투명하게 연결"],
 ["🧪 QA 테스트 대시보드", "핵심 기능 모듈 정량 자동 테스트 결과", "'X/X 통과' — 말이 아닌 숫자로 핵심 기능이 다 돈다 증명"],
 ["🗃️ 데이터 관리(신선도)", "각 데이터가 언제 기준·어디서·자동/수동 갱신인지", "주가 매일 16:30 자동, 재무 공시 때 — 성격에 맞게 설계"],
 ["☁️ AWS 비용·리소스", "배포 인프라 사용량·이번 달 요금", "실제 클라우드 배포 + 비용 통제 증거. GPU는 데모 때만"],
 ["🤖 AI 챗봇 운영(RAG)", "챗봇 서버 실시간 상태·구성·비용(GPU $/시간, 질의당 ~5원)", "임베딩·리랭킹=로컬 GPU(무료), LLM 답변만 유료 구조"],
 ["🖥️ 서비스 생존(Liveness)", "백엔드·프론트·챗봇 포트 연결 + 응답시간", "신호등처럼 생존+속도. 장애 즉시 발견"],
 ["📊 종목 커버리지", "기준 종목 대비 각 기능 커버 %", "주주 100%·재무 99% — 전 종목 보장을 수치로. 갭=NO-MOCK"],
 ["🗄️ DB 현황(16테이블)", "각 테이블 행수(총 약 1,175만 행)", "빈 테이블=빨강(미적재, 정직). 대부분 행은 주가(ohlcv)"],
 ["⏱️ 데이터 신선도", "원천별 가장 최근 갱신 시각(MAX 날짜)", "오래되면 갱신 신호"],
 ["🔎 스모크 테스트", "샘플 종목 핵심 데이터가 실제로 나오는지", "대표 종목 눌러보니 다 나온다 즉석 확인"],
 ["🛠️ 유지보수", "운영자용 관리 동작", "일상 점검·보수용"],
], [1.9, 3.0, 4.0])
quote("운영 탭은 한마디로 '우리 서비스의 건강검진표'입니다. 어떤 기능이 어디서 데이터를 받아 정상인지, 데이터가 며칠 자인지, 비용이 얼마인지 — 매일 5분이면 전체 상태를 파악합니다.")

# ─── 탭2 ───
h("탭 2. 기업분석 · 검증 — '이 데이터가 진짜인지 직접 검사한다'  ★최대 강조")
para("가장 자신 있는 부분입니다. '있다'를 넘어 '맞다'를 증명하는 5단계 검증 — 버튼을 누르면 실제 검사가 돕니다.")
table(["검증", "역할 (무엇을 검사)", "의미·해석"], [
 ["① 내부 정합성 검증", "DB 값 논리 검사 — 재무 항등식(자산=부채+자본)·주가 OHLC·기업개요 완전성·밸류 정합성. 샘플/전수", "숫자끼리 앞뒤가 맞는가. 전수로 전 종목 한 번에 점검"],
 ["② DART 원문 대조", "화면 요약 재무 ↔ DART XBRL 원문 비교", "우리 숫자가 금감원 공시 원본과 일치 — 할루시네이션 아님을 원천으로 증명"],
 ["③ 히스토리브리핑 신뢰도", "AI 브리핑 4-Phase 자동 채점(수치·인용·RAGAS·LLM Judge)", "평균 77.1점·2,526종. AI 요약을 AI가 재검증. 저점수 재생성"],
 ["④ AI 챗봇 검증(RAG)", "챗봇 지식베이스·서버 상태 점검(무료)", "학습 보고서량·검색 생존 확인"],
 ["⑤ 계열사 검증", "계열사 SVG 파일 커버리지·무결성(크기>0·SVG 정상)", "그림 파일 깨짐 점검. GitHub 샘플만→낮은 커버리지 정상"],
], [1.9, 3.3, 3.7])
quote("①은 숫자끼리 논리, ②는 그 숫자가 DART 원문과 같은지, ③은 AI 글 신뢰도, ④챗봇, ⑤그림 — 5개 파트를 원본과 대조합니다. 특히 ② DART 원문 대조가 핵심입니다. 'AI가 지어낸 거 아니냐'는 의심에, 버튼 하나로 공시 원문과 대조해 오류율을 보여드립니다. NO-MOCK을 구호가 아니라 증명으로 만드는 장치입니다.")

# ─── 탭3 ───
h("탭 3. 밸류에이션 · 운영 — '어떤 데이터로, 언제, 어디서 평가했나'  (태상님 매뉴얼 기반)")
para("밸류에이션 투명성 상태판. 적정주가를 만든 재료와 출처를 숨김없이 공개합니다. 네 부분입니다.")
table(["부분", "역할", "의미·해석"], [
 ["① 데이터 기준일", "평가에 쓴 데이터가 며칠 자인지(카드 3장: 밸류·기업개요·챗봇)", "평가일·금리(Rf 4.118%)·시총 세 날짜가 가까운지가 핵심. 벌어지면 재평가. 모델 v9·약 2,231종"],
 ["② 데이터 출처", "쓰는 데이터 출처 8곳 명시", "재무=DART, 주가·시총=KRX, 금리=한국은행(ECOS)… 전부 공식 출처"],
 ["③ 테스트 환경", "엔진 방법(DCF·EPV·멀티플)·챗봇 검색 순서·서버 사양", "기술 질문 대비 사양표"],
 ["④ AI 하네스 ★", "AI가 거짓말 못 하게 막는 4중 장치", "(1)없으면 비움 (2)AI는 '찾기'만·판단은 규칙 (3)모든 숫자에 출처 (4)확신 약하면 숨김(범위만). 핵심 자랑"],
], [1.9, 3.0, 4.0])
quote("'저희가 어떤 데이터를, 언제 기준으로, 어디서 받아 만들었는지'를 직접 확인하시는 곳입니다. 약 2,200종목을 6월 16일 데이터로 평가했고, 금리·회사채·시총 세 날짜를 숨기지 않고 그대로 보여드립니다. 마지막 ④ AI 하네스가 가장 강조하고 싶은 부분 — AI가 숫자를 못 지어내게, 없으면 비우고, 찾기만 맡기고, 출처를 붙이고, 확신 없으면 숨깁니다.")

# ─── 탭4 ───
h("탭 4. 밸류에이션 · 테스트 — '진짜 도는지 직접 눌러본다'  (태상님 매뉴얼 기반)")
para("점검판입니다. 말이 아니라 직접 버튼을 눌러 살아있음을 보여드립니다.")
table(["박스", "역할", "의미·해석"], [
 ["시스템 상태(신호등 3개)", "백엔드(:8090)·챗봇(:8800)·프론트(:3000) 생존", "초록=정상·빨강=응답없음·회색=확인중. 챗봇은 GPU 깨어날 때 잠깐 빨강(정상)"],
 ["🧪 밸류에이션 QA", "종목코드→[밸류 검증]→적정주가·현재가·WACC·신뢰도·EPV + 다섯 체크(✓/✗)", "다섯 ✓=완전 평가. 미평가=‘준비 중’(정상). 대표종목인데 ✗ 많으면 엔진 이상"],
 ["🤖 챗봇 QA", "질문→[질의]→AI 답변 + DART 원문 출처", "핵심은 '답에 출처가 붙느냐'. 출처 없는 단정은 불신. (첫 답변 수십초·~5원)"],
 ["✅ QA 체크리스트", "매번 점검할 5가지", "①없으면 ‘준비 중’? ②일부 빠져도 안 깨짐? ③답에 출처? ④날짜 최신? ⑤신뢰도 등급?"],
], [1.9, 3.2, 3.8])
quote("직접 눌러 확인하는 곳입니다. 신호등 세 개가 초록인 걸 보고, 종목 하나를 검증하면 적정주가·WACC·등급과 함께 꼭 있어야 할 다섯 항목이 전부 체크됩니다. 미평가 종목은 가짜 숫자 대신 '준비 중'으로 정직하게 나옵니다. 저희 원칙은 한 문장 — '실제 데이터만, 출처와 함께, 확신이 없으면 숨긴다.'")

# ─── 마무리 ───
h("마무리 멘트 (관리자 페이지 클로징) — 20초")
quote("일반 화면이 '사용자가 보는 결과'라면, 관리자 페이지는 '저희가 그 결과를 어떻게 책임지는가'입니다. 운영 탭으로 상태를 감시하고, 검증 탭으로 데이터를 원본과 대조하고, 밸류 탭으로 AI의 거짓을 차단합니다. 끝까지 지킨 한 문장으로 마치겠습니다 — '추정하지 않습니다, 출처로 말합니다.'")

# ─── 동선 팁 ───
h("시연 동선 팁 (관리자 파트, 약 3~4분)", 13)
for i, t in enumerate([
 "/admin 접속 → '내부 통제판입니다' (NO-MOCK 한 줄 언급)",
 "운영 탭: 🗺️기능모듈 → 📊커버리지 → 🤖챗봇 비용 (위→아래 드래그하며 '상태판')",
 "검증 탭: ② DART 원문 대조 버튼 실제 클릭(가장 강조) → 5단계 설명",
 "밸류 운영 탭: ④ AI 하네스 강조",
 "밸류 테스트 탭: 신호등 → 밸류 QA 종목 1개 검증 실연 → '준비 중=정직'",
 "클로징: '추정하지 않습니다, 출처로 말합니다.'",
], 1):
    p = doc.add_paragraph(style="List Number"); r = p.add_run(t); setfont(r, 10, False, INK)

doc.save(OUT)
print("저장:", OUT)
