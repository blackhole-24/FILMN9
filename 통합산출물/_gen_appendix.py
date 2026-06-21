# -*- coding: utf-8 -*-
"""
FINSIGHT Appendix (발표 부록) 생성기 — 2026-06-17
전문가 수준(A1~A8) + 초보자 튜토리얼(A9). 무결점 글로벌 서비스 관점.
출력: 통합산출물/FINSIGHT_Appendix_상세.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

NAVY=RGBColor(0x1F,0x3A,0x5F); BLUE=RGBColor(0x25,0x63,0xEB)
GRAY=RGBColor(0x55,0x5F,0x6B); GREEN=RGBColor(0x15,0x7A,0x3C); RED=RGBColor(0xC0,0x2B,0x2B)

def set_kfont(run, name="맑은 고딕"):
    run.font.name=name; rPr=run._element.get_or_add_rPr()
    rFonts=rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts=OxmlElement('w:rFonts'); rPr.append(rFonts)
    for k in ('w:ascii','w:hAnsi','w:eastAsia'): rFonts.set(qn(k),name)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

doc=Document()
st=doc.styles['Normal']; st.font.name='맑은 고딕'; st.font.size=Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.2); s.right_margin=Cm(2.2)

def H(t,size=20,color=NAVY,before=10,after=6,align=None):
    p=doc.add_paragraph(); p.space_before=Pt(before); p.space_after=Pt(after)
    if align is not None: p.alignment=align
    r=p.add_run(t); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color; set_kfont(r); return p

def body(t,size=10,color=None,italic=False,after=4,bold=False,indent=0):
    p=doc.add_paragraph(); p.space_after=Pt(after)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run(t); r.font.size=Pt(size); r.italic=italic; r.bold=bold
    if color: r.font.color.rgb=color
    set_kfont(r); return p

def bullet(t,size=10,color=None,indent=0.6,after=2):
    p=doc.add_paragraph(); p.space_after=Pt(after); p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run("· "+t); r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    set_kfont(r); return p

def chapter(num,title):
    doc.add_paragraph()
    p=doc.add_paragraph(); p.space_before=Pt(10); p.space_after=Pt(8)
    pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'1F3A5F'); pPr.append(sh)
    r=p.add_run(f"  {num}.  {title}"); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); set_kfont(r)

def subh(t):
    p=doc.add_paragraph(); p.space_before=Pt(7); p.space_after=Pt(3)
    r=p.add_run("▣ "+t); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=BLUE; set_kfont(r)

def formula_box(name,lines,variables,note=None):
    p=doc.add_paragraph(); p.space_before=Pt(8); p.space_after=Pt(2)
    r=p.add_run(f"◆ {name}"); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=NAVY; set_kfont(r)
    tbl=doc.add_table(rows=1,cols=1); tbl.style='Table Grid'
    cell=tbl.rows[0].cells[0]; shade(cell,'EAF1FB')
    cp=cell.paragraphs[0]; cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before=Pt(7); cp.paragraph_format.space_after=Pt(7)
    for i,ln in enumerate(lines):
        if i>0: cp.add_run().add_break()
        fr=cp.add_run(ln); fr.bold=True; fr.font.size=Pt(14 if len(lines)==1 else 12.5); fr.font.color.rgb=NAVY; set_kfont(fr)
    ph=doc.add_paragraph(); ph.space_before=Pt(3); ph.space_after=Pt(1); ph.paragraph_format.left_indent=Cm(0.3)
    rh=ph.add_run("▷ 각 값의 뜻"); rh.bold=True; rh.font.size=Pt(9.5); rh.font.color.rgb=GREEN; set_kfont(rh)
    for sym,mean in variables:
        pp=doc.add_paragraph(); pp.space_after=Pt(1); pp.paragraph_format.left_indent=Cm(0.8)
        rs=pp.add_run(sym); rs.bold=True; rs.font.size=Pt(10); rs.font.color.rgb=BLUE; set_kfont(rs)
        rm=pp.add_run("  :  "+mean); rm.font.size=Pt(10); set_kfont(rm)
    if note:
        pn=doc.add_paragraph(); pn.space_after=Pt(4); pn.paragraph_format.left_indent=Cm(0.8)
        rn=pn.add_run("※ "+note); rn.italic=True; rn.font.size=Pt(9.5); rn.font.color.rgb=GRAY; set_kfont(rn)

def table(headers,rows,widths=None,head_fill='1F3A5F'):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,head_fill)
        pr=c.paragraphs[0]; pr.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rr=pr.add_run(h); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); set_kfont(rr)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            pr=cells[i].paragraphs[0]
            rr=pr.add_run(str(val)); rr.font.size=Pt(9); set_kfont(rr)
    if widths:
        for r_ in t.rows:
            for i,w in enumerate(widths): r_.cells[i].width=Cm(w)
    return t

# ───────── 표지 ─────────
H("FINSIGHT",32,NAVY,50,2,WD_ALIGN_PARAGRAPH.CENTER)
H("Appendix · 기술 상세 부록",20,BLUE,2,4,WD_ALIGN_PARAGRAPH.CENTER)
body("발표 후 심화 설명용 — 밸류에이션 계산·RAG·검증체계·인프라·전수검수 + 초보자 튜토리얼",
     11,GRAY,True,4).alignment=WD_ALIGN_PARAGRAPH.CENTER
body("운영사 FILMN9 Inc.  ·  2026-06-17  ·  KPMG AI Lab",10,GRAY,False,2).alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
body("목차",12,NAVY,bold=True)
for x in ["A1. 밸류에이션 정밀 매뉴얼 (6대 공식 + 계산 흐름)",
          "A2. RAG 챗봇 파이프라인 상세","A3. AI 히스토리 브리핑 4-Phase 검증",
          "A4. 데이터 아키텍처 & 4-DB 폴리글랏(ERD)","A5. AWS 인프라 & 무중단·보안 설계",
          "A6. 데이터 원천 지도","A7. 전수 검수 방법론 & 결과","A8. NO-MOCK 신뢰성 검증 체계",
          "A9. 🔰 초보자 튜토리얼 — FINSIGHT 5분 사용법"]:
    body("   "+x,10,GRAY)
doc.add_page_break()

# ═════════ A1 ═════════
chapter("A1","밸류에이션 정밀 매뉴얼")
body("FINSIGHT의 모든 적정주가는 아래 6대 공식으로 계산됩니다. 각 공식과 변수의 뜻을 풀이하고, 전체 계산 흐름을 단계로 정리합니다.",10,GRAY,True)
formula_box("① DCF — 기업가치",
    ["기업가치(EV) = Σ  FCFF(t) ÷ (1 + WACC)^t   +   터미널밸류 ÷ (1 + WACC)^n",
     "적정주가 = ( 기업가치 − 순부채 ) ÷ 발행주식수"],
    [("FCFF(t)","t년차 잉여현금흐름(회사가 자유롭게 쓸 현금)"),("WACC","가중평균자본비용=할인율"),
     ("t","연차(1~n)"),("n","명시적 예측기간(보통 5~10년)"),("터미널밸류","예측기간 이후 영구가치(③)"),
     ("순부채","총차입금 − 보유현금"),("발행주식수","시장 발행 총 주식 수")],
    "미래 현금흐름을 현재가치로 환산·합산 → 기업가치. 채권자 몫을 빼 주주가치, 주식수로 나눠 적정주가.")
formula_box("② FCFF — 잉여현금흐름",
    ["FCFF = EBIT × (1 − Tc) + D&A − CapEx − ΔNWC"],
    [("EBIT","이자·세금 차감 전 영업이익(본업으로 번 돈)"),("Tc","법인세율"),
     ("D&A","감가상각비(현금유출 없는 비용 → +)"),("CapEx","자본적지출(설비투자 현금 → −)"),
     ("ΔNWC","순운전자본 증감(재고·외상에 묶인 돈 증가 → −)")],
    "'장부상 이익'이 아닌 '진짜 남는 현금'. 애널리스트 리포트로 역산·보정해 정밀도를 높임.")
formula_box("③ 터미널밸류 — 영구가치(고든 성장모형)",
    ["터미널밸류 = FCFF(n+1) ÷ (WACC − g) = FCFF(n) × (1 + g) ÷ (WACC − g)"],
    [("FCFF(n+1)","예측 마지막 해 다음 연도 FCFF"),("g","영구성장률(보통 1~3%, 보수적)"),("WACC","할인율")],
    "조건: WACC > g. DCF 가치의 큰 비중 → g 과대설정 금지(NO-MOCK).")
formula_box("④ WACC — 가중평균자본비용",
    ["WACC = (E ÷ V) × Re + (D ÷ V) × Rd × (1 − Tc)"],
    [("E","자기자본(주주 몫, 시가총액)"),("D","타인자본(부채)"),("V","총자본 = E + D"),
     ("Re","자기자본비용(⑤ CAPM)"),("Rd","타인자본비용(차입 이자율)"),("Tc","법인세율(이자 절세효과)")],
    "주주 몫과 빚 몫의 비용을 비중대로 섞은 '종합 자본 단가'.")
formula_box("⑤ CAPM — 자기자본비용",
    ["Re = Rf + β × (Rm − Rf)"],
    [("Re","자기자본비용(구하려는 값)"),("Rf","무위험수익률(국고채 금리)"),
     ("β","베타(시장 대비 민감도)"),("Rm","시장기대수익률"),("(Rm − Rf)","시장위험프리미엄(ERP)")],
    "'안전한 이자 + 위험 보상'으로 주주 요구수익률 계산. β는 피어 추정 후 Blume·Hamada 보정.")
formula_box("⑥ 상대가치 역산 — 4-Way 교차검증",
    ["PER 역산 = 적정 PER × EPS","PBR 역산 = 적정 PBR × BPS",
     "EV/EBITDA : EV = 적정배수 × EBITDA → 적정주가 = (EV − 순부채) ÷ 발행주식수"],
    [("적정 배수","피어 중앙값 등 기준 배수"),("EPS","주당순이익 = 순이익÷주식수"),
     ("BPS","주당순자산 = 순자산÷주식수"),("EBITDA","영업이익+감가상각비(현금창출력 근사)")],
    "DCF(절대) + 3멀티플(상대) = 4-Way 교차검증으로 편향 제거.")
subh("계산 흐름 (엔진 5단계)")
table(["단계","작업","산출"],
    [["1) 재무 수집·표준화","DART XBRL 3개년 파싱 + 신용·발행주식 RAG 추출","정규화 재무"],
     ["2) FCFF 추정","②식 + 애널리스트 리포트 보정","연도별 FCFF"],
     ["3) WACC 산출","④⑤식, 피어 베타 Blume·Hamada 보정","할인율"],
     ["4) DCF 할인·터미널","①③식, Bear/Base/Bull 3시나리오","적정주가 범위"],
     ["5) 4-Way 교차검증·진단","⑥식 + LLM 고/저평가 해설","스코어카드"]],
    widths=[3.5,9,4])
body("한계·정직성: DCF는 가정 기반 '참고치'이므로 단일값이 아닌 범위로 제시하며, 데이터가 없으면 '데이터 준비중'으로 정직 표기(NO-MOCK).",9.5,RED,True,4)

# ═════════ A2 ═════════
chapter("A2","RAG 챗봇 파이프라인 상세")
body("RAG(검색증강생성)는 LLM이 답하기 전 신뢰 가능한 문서를 먼저 검색해 그 근거 위에서 답을 생성하는 구조로, 환각(할루시네이션)을 원천 차단합니다.",10,GRAY,True)
subh("처리 9단계")
table(["#","단계","설명"],
    [["1","문서 수집","전 종목 사업보고서 등 원문 확보"],
     ["2","청크 분할","검색·임베딩에 맞게 일정 크기로 분절"],
     ["3","임베딩","청크를 의미 벡터로 변환"],
     ["4","벡터DB 적재","ChromaDB(약 40GB)에 저장(약 192만 청크 규모)"],
     ["5","질의 임베딩","사용자 질문을 같은 벡터공간으로 변환"],
     ["6","벡터 검색","코사인 유사도로 가까운 청크 후보 추출"],
     ["7","리랭킹","GPU 리랭커로 후보 재정렬(정확도↑, 로컬·무료)"],
     ["8","프롬프트 구성","근거 청크 + 질문 결합"],
     ["9","생성·출처표시","LLM 답변 + 출처 카드·DART 링크 의무 표기"]],
    widths=[1.2,3,12])
body("비용: 임베딩·리랭킹은 로컬 GPU(무료), LLM 답변 호출만 유료(질의당 약 5원 수준). 환각 방지: 근거 없는 답변을 막고 출처를 함께 제시.",9.5,GREEN,True,4)

# ═════════ A3 ═════════
chapter("A3","AI 히스토리 브리핑 4-Phase 검증")
body("사업보고서를 RAG로 요약한 '히스토리 브리핑'의 신뢰도를 4단계로 자동 채점해 등급화합니다.",10,GRAY,True)
table(["Phase","검증","내용"],
    [["1. 수치 매칭","숫자 정확성","브리핑 속 수치가 원본과 일치하는지 대조"],
     ["2. 인용 검증","근거 추적","서술이 실제 문서 근거에 기반하는지 확인"],
     ["3. RAGAS","RAG 품질","검색·생성 적합성 정량 평가"],
     ["4. LLM Judge","종합 심사","LLM이 전체 품질을 종합 채점"]],
    widths=[3,3,10])
body("결과: 약 2,526개 종목 평균 약 77.1점. 저점수 항목은 2-pass 재생성으로 보강. 관리자 신뢰도 검증 탭에서 원본 대조·오류율을 관리.",9.5,GREEN,True,4)

# ═════════ A4 ═════════
chapter("A4","데이터 아키텍처 & 4-DB 폴리글랏 (ERD)")
body("데이터 성격에 따라 4종 DB를 함께 쓰는 폴리글랏 퍼시스턴스 전략을 채택했습니다. (정형·문서·벡터·파일 분리)",10,GRAY,True)
table(["DB","유형","담는 데이터","대표 테이블/객체"],
    [["AWS RDS PostgreSQL","관계형","기업개요·재무·밸류 정형","company_info, financials, ohlcv_daily, valuation_summary, ticker_universe"],
     ["MongoDB Atlas","문서(NoSQL)","AI 히스토리 브리핑(서술형)","histories"],
     ["ChromaDB (EC2)","벡터","사업보고서 임베딩(RAG)","embeddings(약 192만 청크)"],
     ["AWS S3","객체(파일)","Sankey·밸류JSON·계열사 이미지","sankey/, valuation_json/, affiliate_img/"]],
    widths=[3.8,2.2,4.5,6])
subh("ERD 설계 원칙 (정석)")
bullet("관계형 코어(RDS)는 종목코드(stock_code)를 기본키/외래키로 하여 company_info를 루트로 재무·주가·밸류 테이블이 1:N으로 연결.")
bullet("표현 방식: 원형 스파이더웹 ❌ → 가로 방향 트리(좌측 루트 → 우측 가지) + 정석 엔티티-관계(PK/FK, crow's foot) ✅.")
bullet("문서(Mongo)·벡터(Chroma)·파일(S3)은 종목코드를 키로 관계형 코어와 논리적으로 연결(물리적 조인은 아님).")
body("※ 상세 ERD 도식은 별첨 다이어그램(SVG/excalidraw)으로 제공합니다.",9.5,GRAY,True,4)

# ═════════ A5 ═════════
chapter("A5","AWS 인프라 & 무중단·보안 설계")
table(["구성요소","역할","비고"],
    [["EC2 (프론트/백엔드)","Next.js·FastAPI 구동","systemd 자동 기동"],
     ["EC2 GPU (챗봇)","RAG 임베딩·리랭킹·LLM 연계","데모 시에만 가동→정지(비용 최소화)"],
     ["RDS PostgreSQL","정형 데이터(16테이블·약 1,175만 행)","관리형 백업·보안"],
     ["S3","Sankey·밸류·계열사 파일 서빙","대용량 객체"],
     ["MongoDB Atlas","히스토리 브리핑","화이트리스트 접근"],
     ["nginx + HTTPS","리버스프록시·SSL(Let's Encrypt)","/api→백엔드, /→프론트"],
     ["스케줄러","주가 OHLCV 매 거래일 16:00 자동 갱신","무중단 최신화"]],
    widths=[3.8,6.5,5])
subh("가용성·보안 관점 (무결점 서비스 지향)")
bullet("장애 격리: 프론트·백엔드·챗봇·DB를 분리해 한 부분 장애가 전체로 번지지 않음.")
bullet("보안 경계: DB는 보안그룹 화이트리스트로 접근 제한, 비밀키는 .env 런타임 로드(하드코딩·커밋 금지).")
bullet("비용 전략: GPU는 EBS에 데이터 영구 보존 + 인스턴스 정지(stop)로 미사용 시 과금 0, 데모 시 start.")
bullet("HTTPS 전구간 암호화 + 관리자 페이지 토큰 인증.")

# ═════════ A6 ═════════
chapter("A6","데이터 원천 지도")
table(["원천","제공처","수집 데이터","갱신"],
    [["DART OpenAPI","금융감독원","재무제표·사업보고서·공시","실시간/일"],
     ["KRX (pykrx)","한국거래소","일봉 OHLCV","매 거래일 16:00"],
     ["yfinance","글로벌 시세","실시간 시세·글로벌 지수","실시간"],
     ["ECOS","한국은행","국고채 금리 등 거시","일/월"],
     ["WICS","FnGuide","산업분류(업종)","정기"],
     ["네이버","뉴스","종목 뉴스","실시간"]],
    widths=[2.8,3.2,7,2.5])
body("모든 수치·문장은 위 공식 출처에 근거합니다. 출처 없는 데이터는 생성하지 않습니다(NO-MOCK).",9.5,GREEN,True,4)

# ═════════ A7 ═════════
chapter("A7","전수 검수 방법론 & 결과")
body("배포 서비스가 '아무 종목이나 눌러도 깨지지 않는지'를 전 종목·전 기능으로 자동 검증했습니다.",10,GRAY,True)
subh("방법론")
bullet("대상: 활성 2,588종목 × 핵심 11개 API = 28,468건 호출(2026-06-17, AWS 라이브).")
bullet("분류: OK(정상) / EMPTY(데이터 갭·정직표기) / FAIL(5xx·타임아웃·파싱오류=진짜 버그).")
subh("결과")
table(["지표","값"],
    [["총 검사","28,468건 (741초)"],["진짜 버그(FAIL)","8건 → 전수 추적 결과 전부 무해"],
     ["타임아웃 2건","재검사 시 정상(일시적 네트워크)"],["우선주 6건","검색 비노출·도달 불가(무해)"],
     ["실질 결과","데모 차단 버그 0건 (정상률 99.97%)"]],
    widths=[6,10])
body("disclosures(DB)는 비어있으나 화면 전자공시는 라이브 DART(dart_disclosures)·공급계약(supply_contracts)을 사용하므로 무영향. valuation EMPTY는 NO-MOCK 정상(적정가 보유 약 2,223종만 표시).",9.5,GRAY,True,4)

# ═════════ A8 ═════════
chapter("A8","NO-MOCK 신뢰성 검증 체계")
bullet("원칙: 근거 없는 한 글자·숫자도 생성·표시 금지. 데이터가 없으면 '데이터 없음'으로 정직 표기.")
bullet("출처 의무: 모든 수치/문장은 DART·DB·계산식 근거 필수. 챗봇 답변은 출처 카드·DART 링크 동반.")
bullet("관리자 검증 페이지: 전 파트(히스토리·주가/재무·밸류·챗봇) 샘플을 원본과 대조해 오류율을 추적.")
bullet("면책: 본 서비스 정보는 투자 참고용이며, 투자 판단·책임은 이용자 본인에게 있습니다.")

doc.add_page_break()
# ═════════ A9 — 초보자 튜토리얼 ═════════
chapter("A9","🔰 초보자 튜토리얼 — FINSIGHT 5분 사용법")
body("주식·재무를 처음 접하는 분도 따라 할 수 있도록, 화면 순서대로 설명합니다. 어려운 단어는 본 부록 뒤의 [용어사전]을 참고하세요.",10,GRAY,True)
subh("STEP 1. 메인 화면 — 원하는 회사 찾기")
bullet("상단 검색창에 회사 이름(예: 삼성전자)이나 업종(예: 반도체)을 입력하면 결과가 그룹으로 나옵니다.")
bullet("가운데 '추천 종목'이 자동으로 돌아가고, 오른쪽 '글로벌 마켓 시그널'에서 오늘 시장 분위기(다우·나스닥·환율·공포탐욕)를 봅니다.")
bullet("관심 있는 종목은 ★(즐겨찾기)로 저장하면 '내 관심종목'에 모입니다.")
subh("STEP 2. 종목 상세 — 3개 탭으로 분석")
table(["탭","무엇을 보나","쉽게 말하면"],
    [["기업개요","회사정보·재무 하이라이트·전자공시·계열사·손익흐름도","이 회사가 뭐 하는 회사이고 돈을 잘 버는지"],
     ["밸류에이션","적정주가(DCF 3시나리오)·4-Way·스코어카드","지금 주가가 싼지 비싼지"],
     ["AI 챗봇","질문하면 사업보고서 근거로 답변+출처","궁금한 걸 물어보는 상담원"]],
    widths=[3,8,6])
subh("STEP 3. 숫자 읽는 법 (초간단)")
bullet("재무 하이라이트의 'YoY'는 작년 같은 때보다 얼마나 늘었나(+)/줄었나(−) 입니다.")
bullet("밸류에이션의 '상승여력 +20%'는 적정가가 현재가보다 20% 높다는 뜻(저평가 신호).")
bullet("'데이터 준비중/없음'은 거짓 숫자를 만들지 않기 위해 정직하게 비워둔 것입니다(NO-MOCK).")
subh("STEP 4. AI 챗봇에 물어보기")
bullet("예: '삼성전자 최근 실적 어때?' → 사업보고서를 찾아 요약하고, 아래에 근거 출처와 DART 링크를 보여줍니다.")
bullet("출처가 함께 나오므로, 답이 어디서 왔는지 직접 확인할 수 있습니다.")
subh("💡 화면 속 (i) 도움말")
body("어려운 용어 옆의 (i) 아이콘에 마우스를 올리면 한 줄 설명이 뜹니다. 더 자세한 건 [용어사전]에서 찾아보세요.",10,BLUE,False,4)

doc.add_paragraph()
body("— 본 부록의 정의·수치는 공식 출처와 계산식에 근거하며, 투자 판단의 책임은 이용자에게 있습니다. —",8.5,GRAY,True)

out=Path(r"C:\Users\Admin\FILMN9\통합산출물\FINSIGHT_Appendix_상세.docx")
doc.save(out)
print("저장:",out)
print("표:",len(doc.tables),"개 / 문단:",len(doc.paragraphs),"개")
